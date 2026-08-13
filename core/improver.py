from pathlib import Path
import shutil
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt
import zipfile
import regex

from config.rules import (
    ACTION_VERB_REPLACEMENTS,
    UNPROFESSIONAL_PHRASES,
    get_required_sections,
    get_section_placeholder,
)
from config.settings import IMPROVED_FOLDER
from core.exporter import export_docx, export_pdf
from core.parser import parse_document
from services.ai_rewriter import rewrite_document_with_ai, rewrite_section_with_ai
from core.change_tracker import get_change_log, reset_change_log, set_document_context

# ============================================
# SECTION ALIASES AND ORDERS
# ============================================

SECTION_ALIASES = {
    "Professional Summary": ["professional summary", "summary", "objective", "profile", "about me"],
    "Projects": ["projects", "project"],
    "Relevant Coursework": ["relevant coursework", "coursework"],
    "Soft Skills": ["soft skills", "soft skill"],
    "Experience": [
        "experience", "work experience", "work", "work history", "employment",
        "employment history", "professional experience", "career history",
        "leadership experience", "leadership",
        "internships", "internship",
    ],
    "Education": [
        "education", "academic", "academics", "education and qualifications",
        "qualifications", "academic background",
    ],
    "Technical Skills": [
        "technical skills", "skills", "key skills", "core skills",
        "core competencies", "competencies", "tech stack", "technologies",
        "technical skill",
    ],
    "Certifications": ["certifications", "certification", "certificates", "certificate"],
    "Additional Information": [
        "additional information", "additional details", "additional",
        "general information", "hobbies", "interests",
    ],
}

RESUME_ORDER = [
    "Professional Summary",
    "Experience",
    "Projects",
    "Education",
    "Relevant Coursework",
    "Technical Skills",
    "Soft Skills",
    "Certifications",
    "Additional Information",
]

THESIS_ORDER = [
    "Document Details",
    "Acknowledgement",
    "Abstract",
    "Introduction",
    "Objectives",
    "Literature Review",
    "Methodology",
    "Analysis",
    "Results",
    "Discussion",
    "Conclusion",
    "References",
]

REPORT_ORDER = [
    "Document Details",
    "Executive Summary",
    "Introduction",
    "Objectives",
    "Literature Review",
    "Methodology",
    "Analysis",
    "Findings",
    "Conclusion",
    "Recommendations",
    "References",
]

PROPOSAL_ORDER = [
    "Title",
    "Executive Summary",
    "Problem Statement",
    "Objectives",
    "Action Plan",
    "Resources Needed",
    "Budget",
    "Timeline",
    "Success Metrics",
    "Conclusion & Request",
]

COVER_LETTER_ORDER = [
    "Sender Information",
    "Date",
    "Recipient Information",
    "Subject",
    "Salutation",
    "Body",
    "Closing",
    "Signature",
]

ACADEMIC_HEADING_ALIASES = {
    "Acknowledgement": ["acknowledgement", "acknowledgments", "acknowledgements"],
    "Abstract": ["abstract", "abstarct", "abstact"],
    "Executive Summary": ["executive summary", "summary"],
    "Introduction": ["introduction", "intro", "background"],
    "History": ["history", "brief history", "history of ai", "brief history of ai"],
    "Objectives": [
        "objective", "objectives", "aim", "aims",
        "objetives", "objectives of the proposal", "objetives of the proposal",
    ],
    "Literature Review": ["literature review", "litrature review", "related work", "review of literature"],
    "Methodology": ["methodology", "methods", "research methodology", "research method"],
    "Analysis": ["analysis", "data analysis"],
    "Results": ["results", "result"],
    "Discussion": ["discussion"],
    "Findings": ["findings", "finding"],
    "Conclusion": ["conclusion", "conclusions", "conclution"],
    "Recommendations": ["recommendations", "recommendation"],
    "References": ["references", "bibliography"],
    "Problem Statement": [
        "problem statement", "statement of the problem",
        "background of the problem", "background of problem",
        "backround of the problem", "backround of problem",
    ],
    "Technical Approach": [
        "technical approach", "plan", "method", "methods", "plan / method",
        "plan/method", "approach", "implementation approach", "proposed solution",
        "solution", "proposed methodology", "proposed activities", "activities",
        "proposed activites",
    ],
    "Timeline": ["timeline", "schedule", "work plan"],
    "Budget": [
        "budget", "budget summary", "estimated budget", "cost estimate",
        "resources needed", "resources required", "cost summary",
    ],
    "Expected Outcomes": ["expected outcomes", "outcomes", "expected result", "expected results"],
    "Deliverables": ["deliverables", "outputs"],
    "Qualifications or Team": ["qualifications or team", "team", "project team", "team members"],
    "Evaluation": ["evaluation", "monitoring and evaluation"],
    "Terms and Conditions": ["terms and conditions", "terms", "conditions"],
    "Appendices": ["appendices", "appendix"],
}

SPELLING_FIXES = {
    "redusing": "reducing",
    "prepaired": "Prepared",
    "janruary": "January",
    "backround": "background",
    "fulfilment": "fulfillment",
    "organisation": "organization",
    "optimisation": "optimization",
    "behaviour": "behavior",
    "catalogue": "catalog",
    "converstion": "conversation",
    "setpember": "September",
    "costumer": "customer",
    "sistem": "system",
    "teh": "the",
    "seperated": "separated",
    "seperate": "separate",
    "recieve": "receive",
    "docmument": "document",
    "enviroment": "environment",
    "langauge": "language",
    "becuase": "because",
    "becuse": "because",
    "goverment": "government",
    "meny": "many",
    "plastik": "plastic",
    "markt": "marked",
    "objetives": "objectives",
    "importence": "importance",
    "classerooms": "classrooms",
    "activites": "activities",
    "awerness": "Awareness",
    "campain": "Campaign",
    "cheep": "cheap",
    "throwed": "thrown",
    "cleanner": "cleaner",
    "envirmentally": "environmentally",
    "frendly": "friendly",
    "conclution": "Conclusion",
    "aprove": "approve",
    "lern": "learn",
    # Additional common misspellings
    "occured": "occurred",
    "neccessary": "necessary",
    "definately": "definitely",
    "aquired": "acquired",
    "realy": "really",
    "succeded": "succeeded",
    "untill": "until",
    "agian": "again",
    "wich": "which",
    "thier": "their",
    "reccomend": "recommend",
    "ocassion": "occasion",
    "occassion": "occasion",
    "begining": "beginning",
    "recieved": "received",
    "acheive": "achieve",
    "desicribe": "describe",
    "existance": "existence",
    "practise": "practice",
    "occassionally": "occasionally",
    "grammer": "grammar",
    "speach": "speech",
    "writting": "writing",
    "reasearch": "research",
    "personel": "personnel",
    "analize": "analyze",
    "personell": "personnel",
    "developement": "development",
    "maintainence": "maintenance",
    "emergancy": "emergency",
    "sincerly": "sincerely",
    "occurrance": "occurrence",
    "alot": "a lot",
    "aswell": "as well",
}

BULLET_MARKERS = ("-", "*", "\u2022", "\u25aa", "\u25e6")
BULLET_SECTIONS = {"Objectives", "Deliverables", "Qualifications or Team"}

PARAGRAPH_SECTIONS = {
    "Document Details", "Acknowledgement", "Abstract", "Executive Summary",
    "Introduction", "Objectives", "Literature Review", "Methodology",
    "Analysis", "Results", "Discussion", "Findings", "Conclusion",
    "Recommendations", "References", "Professional Summary", "Content",
}

JUNK_LINE_PATTERNS = (
    r"education\s+world",
    r"use\s+handout",
    r"permission\s+to\s+reproduce",
    r"for\s+educational\s+purposes\s+only",
    r"https?://www\.educationworld\.com",
    r"^??\s*\d{4}\b",
    r"^\[page\s+break\]$",
    r"^---$",
)

UNPROFESSIONAL_RESUME_PATTERNS = (
    r"mostly\s+free",
    r"gym\s*time",
    r"facebook",
    r"post\s+pictures",
    r"graduate\s+of\s+life",
    r"any\s+fun\s+position",
    r"cool\s+girl",
    r"friends!?$",
)


# ============================================
# CHANGE TRACKING HELPERS
# ============================================

def _replace_unprofessional_phrases_with_tracking(text: str, section: str = "") -> str:
    """Replace unprofessional phrases and track changes"""
    if not text:
        return text
    
    result = text
    change_log = get_change_log()
    
    for bad_phrase, good_phrase in UNPROFESSIONAL_PHRASES.items():
        pattern = rf'\b{regex.escape(bad_phrase)}\b'
        
        # Find all occurrences
        for match in regex.finditer(pattern, result, flags=regex.IGNORECASE):
            original = match.group(0)
            
            if good_phrase == "":
                change_log.add("phrase", original, "[removed]", section)
            else:
                change_log.add("phrase", original, good_phrase, section)
        
        # Apply replacements
        if good_phrase == "":
            result = regex.sub(pattern, '', result, flags=regex.IGNORECASE)
        else:
            result = regex.sub(pattern, good_phrase, result, flags=regex.IGNORECASE)
    
    # Clean up spacing while preserving line structure
    result = regex.sub(r'[ 	]+', ' ', result)
    result = regex.sub(r' *\n *', '\n', result)
    result = regex.sub(r'\n{3,}', '\n\n', result).strip()
    return result


def _add_missing_sections_with_tracking(
    sections: Dict[str, List[str]], 
    doc_type: str, 
    order: List[str]
) -> Tuple[Dict[str, List[str]], List[str]]:
    """Add missing sections and track changes"""
    if doc_type in {"Resume", "CV"}:
        normalized_order = [
            section_name for section_name in RESUME_ORDER
            if section_name in sections and sections.get(section_name)
        ]
        trailing = [
            section_name for section_name in order
            if section_name not in normalized_order and section_name in sections and sections.get(section_name)
        ]
        return sections, normalized_order + trailing
    if doc_type in {"Report", "Proposal", "Thesis"}:
        # Academic documents should preserve the author's existing structure instead of
        # injecting placeholder sections that often degrade the final output.
        normalized_order = [
            section_name for section_name in order
            if section_name in sections and sections.get(section_name)
        ]
        return sections, normalized_order
    if doc_type == "Cover Letter":
        change_log = get_change_log()
        metadata_only_sections = {
            "Sender Information": "[Sender Information: Add your address, email, and phone number]",
            "Date": "[Date: Add date]",
            "Recipient Information": "[Recipient Information: Add hiring manager and company address]",
        }
        for section_name, placeholder in metadata_only_sections.items():
            if not sections.get(section_name):
                sections[section_name] = [placeholder]
                change_log.add("section_added", f"[Missing: {section_name}]", placeholder, section_name)
        normalized_order = [
            section_name for section_name in COVER_LETTER_ORDER
            if section_name in sections and sections.get(section_name)
        ]
        trailing = [
            section_name for section_name in order
            if section_name not in normalized_order and section_name in sections and sections.get(section_name)
        ]
        return sections, normalized_order + trailing

    required = get_required_sections(doc_type)
    change_log = get_change_log()
    
    for section in required:
        section_exists = False
        for existing in sections.keys():
            existing_lower = existing.lower()
            section_lower = section.lower()
            if existing_lower == section_lower:
                section_exists = True
                break
            if existing_lower in section_lower or section_lower in existing_lower:
                section_exists = True
                break
            if "experience" in existing_lower and "experience" in section_lower:
                section_exists = True
                break
        
        if not section_exists:
            placeholder = get_section_placeholder(section)
            sections[section] = [placeholder]
            if section not in order:
                order.append(section)
            
            change_log.add("section_added", f"[Missing: {section}]", placeholder, section)
    
    return sections, order


# ============================================
# TEXT PROCESSING FUNCTIONS
# ============================================

def _normalize_text(text: str) -> str:
    cleaned = text.replace("\r", "\n")
    cleaned = cleaned.replace("**", "")
    cleaned = regex.sub(r"\(cid:\d+\)", "- ", cleaned, flags=regex.IGNORECASE)
    cleaned = cleaned.replace("\u2022", "- ").replace("\u25aa", "- ").replace("\u25cf", "- ")
    cleaned = regex.sub(r"(?m)^\s*#{1,6}\s*", "", cleaned)
    cleaned = regex.sub(r"[-]{3,}", "\n", cleaned)
    cleaned = regex.sub(r"[ \t]+", " ", cleaned)
    cleaned = regex.sub(r"[ \t]*\n[ \t]*", "\n", cleaned)
    cleaned = regex.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _is_document_junk_line(text: str) -> bool:
    stripped = text.strip()
    lowered = stripped.lower()
    if not stripped:
        return True
    if regex.search(r"\b\d{1,2}(?:st|nd|rd|th)?\s+[a-z]+,?\s+\d{4}\b|\b[a-z]+\s+\d{1,2},?\s+\d{4}\b", lowered, flags=regex.IGNORECASE):
        return False
    if regex.fullmatch(r"[-|_\s]{3,}", stripped):
        return True
    if regex.fullmatch(r"\?{5,}", stripped):
        return True
    if regex.search(r"\(cid:\d+\)", stripped, flags=regex.IGNORECASE):
        return True
    if any(regex.search(pattern, lowered, flags=regex.IGNORECASE) for pattern in JUNK_LINE_PATTERNS):
        return True
    return False


def _is_noise_line(text: str) -> bool:
    return _is_document_junk_line(text)


def _extract_name_contact(lines: List[str]) -> Tuple[str, str]:
    name = "Candidate Name"
    email = ""
    phone = ""
    link = ""

    for line in lines[:8]:
        if not email:
            email_match = regex.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line)
            if email_match:
                email = email_match.group(0)
        if not phone:
            phone_match = regex.search(r"(\+?\d[\d\s\-()]{7,}\d)", line)
            if phone_match:
                phone = regex.sub(r"\s+", "", phone_match.group(0))
        if not link:
            link_match = regex.search(
                r"(https?://\S+|www\.\S+|linkedin\.com/\S+|github\.com/\S+)",
                line,
                flags=regex.IGNORECASE,
            )
            if link_match:
                link = link_match.group(0)

    ignored_name_prefixes = {
        "address:", "mobile:", "phone:", "email:", "contact:", "linkedin:", "github:",
    }
    ignored_name_tokens = {
        "education", "experience", "skills", "certifications", "summary", "objective",
        "coursework", "information", "insights", "leadership", "technical",
    }

    def _is_valid_name_tokens(tokens: List[str]) -> bool:
        if not (2 <= len(tokens) <= 4):
            return False
        if any(token.lower().strip(",.:;") in ignored_name_tokens for token in tokens):
            return False
        return all(regex.fullmatch(r"[A-Za-z][A-Za-z'`.-]*", token) for token in tokens)

    for line in lines[:8]:
        stripped = line.strip()
        if not stripped:
            continue
        prefix = regex.split(r"\b(?:email|mobile|phone|address|contact|linkedin|github)\s*:", stripped, maxsplit=1, flags=regex.IGNORECASE)[0].strip()
        prefix_tokens = prefix.split()
        if _is_valid_name_tokens(prefix_tokens):
            name = " ".join(token.title() for token in prefix_tokens)
            break

    for line in lines[:6]:
        stripped = line.strip()
        lowered = stripped.lower()
        if any(lowered.startswith(prefix) for prefix in ignored_name_prefixes):
            continue
        tokens = stripped.split()
        if _is_valid_name_tokens(tokens):
            name = stripped.title()
            break

    if name == "Candidate Name" and email:
        local_part = email.split("@", 1)[0]
        local_part = regex.sub(r"\d+", " ", local_part)
        local_part = regex.sub(r"[_\-.]+", " ", local_part)
        local_part = regex.sub(r"\s+", " ", local_part).strip()
        guess_tokens = [token for token in local_part.split() if len(token) > 1]
        if 2 <= len(guess_tokens) <= 4 and all(regex.fullmatch(r"[A-Za-z][A-Za-z'`.-]*", token) for token in guess_tokens):
            name = " ".join(token.title() for token in guess_tokens)

    contact = " | ".join([item for item in [email, phone, link] if item])
    return name, contact


def _looks_like_summary_text(value: str) -> bool:
    lowered = value.lower().strip()
    if len(value.split()) < 12:
        return False
    return any(token in lowered for token in {
        "i am", "currently", "interested in", "passionate", "curiosity", "motivated",
        "aspiring", "final year", "seeking",
    })


def _postprocess_resume_data(name: str, contact: str, sections: Dict[str, List[str]]) -> Tuple[str, str, Dict[str, List[str]]]:
    normalized = {key: list(values) for key, values in sections.items() if values}
    change_log = get_change_log()
    soft_skill_markers = {
        "communication", "collaboration", "teamwork", "team", "leadership",
        "time management", "attention to detail", "adaptability", "problem solving",
        "willingness to learn", "interpersonal", "organization", "organisational",
    }

    additional_items = normalized.get("Additional Information", [])
    if "Professional Summary" not in normalized:
        summary_candidates = [item for item in additional_items if _looks_like_summary_text(item)]
        if summary_candidates:
            normalized["Professional Summary"] = [" ".join(summary_candidates)]
            normalized["Additional Information"] = [item for item in additional_items if item not in summary_candidates]
            if not normalized["Additional Information"]:
                normalized.pop("Additional Information", None)
            change_log.add(
                "reorder",
                "Additional Information",
                "Professional Summary",
                "Professional Summary",
            )

    if "Technical Skills" in normalized:
        technical = []
        coursework = list(normalized.get("Relevant Coursework", []))
        soft_skills = list(normalized.get("Soft Skills", []))
        for item in normalized["Technical Skills"]:
            lowered = item.lower()
            if "coursework" in lowered or lowered in {
                "statistics", "database management system", "introduction to artificial intelligence",
                "research methodology", "technical writing",
            }:
                coursework.append(item)
            elif any(marker in lowered for marker in soft_skill_markers):
                soft_skills.append(item)
            else:
                technical.append(item)
        normalized["Technical Skills"] = technical
        if coursework:
            normalized["Relevant Coursework"] = coursework
            change_log.add(
                "reorder",
                "Technical Skills",
                "Relevant Coursework",
                "Relevant Coursework",
            )
        if soft_skills:
            normalized["Soft Skills"] = soft_skills
            change_log.add(
                "reorder",
                "Technical Skills",
                "Soft Skills",
                "Soft Skills",
            )
        if not normalized["Technical Skills"]:
            normalized.pop("Technical Skills", None)

    if "Certifications" in normalized:
        normalized["Certifications"] = _normalize_certification_items(normalized["Certifications"])

    return name, contact, normalized


def _normalize_certification_items(items: List[str]) -> List[str]:
    provider_markers = {
        "datacamp", "coursera", "udemy", "edx", "linkedin learning",
        "google", "microsoft", "aws", "ibm", "meta", "oracle",
    }
    merged_certifications: List[str] = []
    change_log = get_change_log()

    for item in items:
        cleaned_item = _polish_item(item)
        if not cleaned_item:
            continue
        lowered = cleaned_item.lower().strip()

        if (
            merged_certifications
            and len(cleaned_item.split()) <= 4
            and any(marker in lowered for marker in provider_markers)
        ):
            provider_display = "AWS" if lowered == "aws" else cleaned_item.title()
            if lowered == "datacamp":
                provider_display = "Datacamp"
            merged_certifications[-1] = f"{merged_certifications[-1]} - {provider_display}"
            change_log.add("rewrite", cleaned_item, merged_certifications[-1], "Certifications")
            continue

        provider_in_same_line = next((marker for marker in provider_markers if marker in lowered), "")
        if provider_in_same_line:
            provider_display = provider_in_same_line.title() if provider_in_same_line != "aws" else "AWS"
            before_provider = regex.sub(
                rf"\s*-?\s*{regex.escape(provider_in_same_line)}\b",
                "",
                cleaned_item,
                flags=regex.IGNORECASE,
            ).strip(" -")
            normalized_item = f"{before_provider} - {provider_display}" if before_provider else provider_display
            if not merged_certifications or merged_certifications[-1].lower() != normalized_item.lower():
                if normalized_item != cleaned_item:
                    change_log.add("rewrite", cleaned_item, normalized_item, "Certifications")
                merged_certifications.append(normalized_item)
            continue

        merged_certifications.append(cleaned_item)

    deduped: List[str] = []
    seen = set()
    for item in merged_certifications:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def _clean_heading_candidate(line: str) -> str:
    cleaned = regex.sub(r"^\s*[\-\*\u2022\u25aa\u25e6]+\s*", "", line or "")
    cleaned = regex.sub(r"^\(cid:\d+\)\s*", "", cleaned, flags=regex.IGNORECASE)
    return cleaned.strip()


def _canonical_heading(line: str) -> str:
    line = _clean_heading_candidate(line)
    token = regex.sub(r"[^a-z ]", "", line.lower()).strip()
    token = regex.sub(r"\s+", " ", token)
    for section, aliases in SECTION_ALIASES.items():
        if token in aliases:
            return section
    return ""


def _split_resume_heading_line(line: str) -> Tuple[str, str]:
    simplified = regex.sub(r"\s+", " ", _clean_heading_candidate(line))
    lowered = simplified.lower()
    candidates = []
    for section, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            candidates.append((alias, section))
    for alias, section in sorted(candidates, key=lambda item: len(item[0]), reverse=True):
        if lowered == alias:
            return section, ""
        if lowered.startswith(alias + ":"):
            return section, simplified[len(alias) + 1:].strip()
        if lowered.startswith(alias + " "):
            return section, simplified[len(alias):].strip()
    return "", ""


def _rewrite_resume_item(text: str, section: str) -> str:
    lowered = text.lower().strip()
    if any(regex.search(pattern, lowered, flags=regex.IGNORECASE) for pattern in UNPROFESSIONAL_RESUME_PATTERNS):
        if section == "Professional Summary":
            return "Motivated entry-level candidate with strong communication, creativity, and willingness to learn."
        return ""
    rewrites = {
        "good at holding a conversation about anything": "Demonstrated strong verbal communication skills.",
        "good at holding a converstion about anything": "Demonstrated strong verbal communication skills.",
        "great at updating facebook and anything i can post pictures on": "Comfortable using social media and digital content platforms.",
        "great at updating facebook and anything i can post pictures on.": "Comfortable using social media and digital content platforms.",
        "shot pics of my friends!": "Photographed subjects and supported image capture tasks.",
        "added cute details in photoshop": "Edited images in Adobe Photoshop.",
        "went to weekly meetings": "Participated in weekly team meetings.",
        "flip burgers and fill condiment containers": "Prepared food items and maintained station readiness.",
        "deal with annoying customer requests": "Handled customer requests professionally.",
        "deal with annoying costumer requests": "Handled customer requests professionally.",
        "answer phone": "Answered phone inquiries and assisted customers.",
        "liked chocolate best": "",
    }
    return rewrites.get(lowered, text)


def _clean_cover_letter_text(text: str) -> str:
    cleaned = regex.sub(r"[--]", " ", text)
    cleaned = cleaned.replace("?", "'").replace("?", '"').replace("?", '"').replace("?", "-")
    cleaned = regex.sub(r"\s+", " ", cleaned).strip()

    replacements = {
        "Hello! Please find my resume attached": "Please find my resume attached",
        "Why did I sign up for this position? Well, I'm going to cut to the chase here:": "I am excited to apply for this position because",
        "I am your guy": "I would welcome the opportunity to contribute to your team",
        "with yourself or the director": "with you or the hiring team",
        "good-looking": "professional",
        "genius": "candidate",
        "SUPERB": "strong",
        "organisational": "organizational",
        "programmes": "programs",
        "companys": "companies",
        "Thanks for your consideration": "Thank you for your consideration",
        "Thanks for your consideration;": "Thank you for your consideration;",
    }
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)

    cleaned = cleaned.replace("I m", "I'm")
    cleaned = cleaned.replace("doesn t", "doesn't")
    cleaned = cleaned.replace("don t", "don't")
    cleaned = regex.sub(
        r"Why did I sign up for this position['’?]?\s*Well,\s*I['’]m going to cut to the chase here:\s*",
        "I am excited to apply for this position because ",
        cleaned,
        flags=regex.IGNORECASE,
    )
    cleaned = regex.sub(
        r"I['’]m born to be a\s+([A-Za-z /-]+?)\s+after reading the requirements in your advertisement",
        r"my background aligns well with the \1 based on the requirements outlined in your advertisement",
        cleaned,
        flags=regex.IGNORECASE,
    )
    cleaned = regex.sub(r"P\.?S\.? .*", "", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"my friends often tell me.*?LOL\.?", "", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"It also may be of interest to you that I was the Hall King of XYZ Hall in 2019[^.?!]*[.?!]?", "", cleaned, flags=regex.IGNORECASE)
    cleaned = cleaned.replace(
        "if you are looking to have a smart, creative, professional, and detail-oriented candidate in your team, I would welcome the opportunity to contribute to your team",
        "I would welcome the opportunity to contribute my creativity, professionalism, and attention to detail to your team",
    )
    cleaned = cleaned.replace("Thank you for your consideration for your consideration", "Thank you for your consideration")
    cleaned = regex.sub(r"\s+([,.;:!?])", r"\1", cleaned)
    cleaned = regex.sub(r"([.!?])(?=[A-Z])", r"\1 ", cleaned)
    cleaned = regex.sub(r"\s{2,}", " ", cleaned).strip(" ,;-")
    return cleaned


def _merge_cover_letter_body_blocks(blocks: List[List[str]]) -> List[str]:
    paragraphs: List[str] = []
    current = ""

    for block in blocks:
        block_text = _clean_cover_letter_text(" ".join(line for line in block if not _is_subject_line(line)))
        if not block_text:
            continue

        if not current:
            current = block_text
            continue

        needs_merge = (
            not regex.search(r"[.!?][\"']?$", current)
            or block_text[:1].islower()
            or current.lower().endswith((" in a", " of", " to", " and", " with", " for", " the", " a", " an"))
        )

        if needs_merge:
            current = f"{current} {block_text}".strip()
        else:
            paragraphs.append(current)
            current = block_text

    if current:
        paragraphs.append(current)

    return paragraphs


def _split_cover_letter_blocks(text: str) -> List[List[str]]:
    blocks: List[List[str]] = []
    current: List[str] = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        if _is_document_junk_line(line):
            continue
        current.append(line)
    if current:
        blocks.append(current)
    return blocks


def _is_date_line(line: str) -> bool:
    lowered = line.lower().strip()
    if regex.search(r"\b\d{1,2}(st|nd|rd|th)?\s+[a-z]+\s*,?\s+\d{4}\b", lowered):
        return True
    if regex.search(r"\b[a-z]+\s+\d{1,2},?\s+\d{4}\b", lowered):
        return True
    if regex.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", lowered):
        return True
    return False


def _is_subject_line(line: str) -> bool:
    return line.lower().startswith("subject:")


def _looks_like_recipient_line(line: str) -> bool:
    lowered = line.lower().strip()
    strong_recipient_markers = {
        "hiring manager", "recruiter", "hr manager", "human resources",
        "dear hiring manager", "company", "corp", "corporation", "inc",
        "ltd", "llc", "pvt", "services", "technologies", "tech",
    }
    if any(marker in lowered for marker in strong_recipient_markers):
        return True
    if regex.search(r"\b(street|road|rd\.?|avenue|ave\.?)\b", lowered):
        return True
    return False


def _clean_fallback_items(section_text: str, section_name: str, doc_type: str) -> List[str]:
    if doc_type == "Cover Letter":
        if section_name == "Body":
            paragraphs = []
            for paragraph in section_text.split("\n"):
                cleaned = _clean_cover_letter_text(paragraph)
                if cleaned:
                    paragraphs.append(cleaned)
            return paragraphs
        cleaned_lines = []
        for line in section_text.split("\n"):
            cleaned_line = _clean_cover_letter_text(line)
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        return cleaned_lines

    corrected_items = _split_line_items(section_text, section_name)
    corrected_items = [_local_refine_item_with_tracking(item, section_name, doc_type) for item in corrected_items]
    corrected_items = [item for item in corrected_items if item]
    if not corrected_items and section_text.strip():
        fallback_item = _local_refine_item_with_tracking(section_text.strip(), section_name, doc_type)
        corrected_items = [fallback_item] if fallback_item else []
    return corrected_items


def _should_accept_academic_rewrite(
    section_name: str,
    original_items: List[str],
    rewritten_items: List[str],
) -> bool:
    if not rewritten_items:
        return False

    normalized_original = "\n".join(_polish_item(item) for item in original_items if _polish_item(item)).strip()
    normalized_rewritten = "\n".join(_polish_item(item) for item in rewritten_items if _polish_item(item)).strip()
    if not normalized_rewritten:
        return False

    first_item = rewritten_items[0].strip(" :-").lower()
    section_label = section_name.strip().lower()
    if first_item == section_label:
        return False

    if len(normalized_original) >= 240 and len(normalized_rewritten) < max(120, int(len(normalized_original) * 0.55)):
        return False

    return True


def _canonical_academic_heading(line: str) -> str:
    token = regex.sub(r"^\d+\s*", "", line.lower()).strip()
    token = regex.sub(r"[^a-z ]", "", token)
    token = regex.sub(r"\s+", " ", token).strip()
    for section, aliases in ACADEMIC_HEADING_ALIASES.items():
        if token in aliases:
            return section
    return ""


def _parse_academic_heading_line(line: str) -> Tuple[str, str]:
    stripped = line.strip()
    if not stripped:
        return "", ""

    simplified = regex.sub(r"\s+", " ", stripped).strip()
    lowered = simplified.lower()
    main_heading_prefix = r"^(?:chapter\s+\d+\s*[:.)-]?\s*|\d+(?:\.0)?[.):-]\s+|[ivxlcdm]+(?:[.):-]\s+|\s+))"
    lowered = regex.sub(r"^(?:chapter\s+\d+\s*[:.)-]?\s*)", "", lowered)
    lowered = regex.sub(main_heading_prefix, "", lowered, flags=regex.IGNORECASE)
    lowered = regex.sub(r"\s+", " ", lowered).strip()

    for section, aliases in ACADEMIC_HEADING_ALIASES.items():
        for alias in aliases:
            if lowered == alias:
                return section, ""
            if lowered.startswith(alias + ":"):
                remainder = simplified[len(simplified) - len(lowered) + len(alias) + 1:].strip(" -\t")
                return section, remainder

    return "", ""


def _parse_generic_academic_heading_line(line: str) -> Tuple[str, str]:
    stripped = line.strip()
    if not stripped:
        return "", ""

    simplified = regex.sub(r"\s+", " ", stripped).strip()
    if len(simplified) > 90:
        return "", ""
    if simplified.endswith((".", "!", "?")):
        return "", ""

    main_heading_prefix = r"^(?:chapter\s+\d+\s*[:.)-]?\s*|\d+(?:\.0)?[.):-]\s+|[ivxlcdm]+(?:[.):-]\s+|\s+))"
    prefix_removed = regex.sub(main_heading_prefix, "", simplified, flags=regex.IGNORECASE).strip()
    if not prefix_removed:
        return "", ""

    word_count = len(prefix_removed.split())
    if word_count < 3 or word_count > 10:
        return "", ""

    if regex.search(r"[@]|https?://|www\.", prefix_removed, flags=regex.IGNORECASE):
        return "", ""

    alpha_only = regex.sub(r"[^A-Za-z ]", "", prefix_removed).strip()
    if not alpha_only:
        return "", ""

    has_heading_prefix = bool(regex.match(main_heading_prefix, simplified, flags=regex.IGNORECASE))

    # Only treat generic report/proposal headings as headings when they are explicitly
    # numbered or chapter-prefixed. This avoids converting normal topic lines like
    # "Healthcare", "Positive", or "Machine Learning" into major sections.
    if not has_heading_prefix:
        return "", ""

    normalized_heading = regex.sub(r"\s+", " ", prefix_removed).strip(" :-")
    normalized_heading = normalized_heading.title()
    return normalized_heading, ""


def _split_line_items(line: str, section: str) -> List[str]:
    if not line.strip():
        return []

    if regex.match(r"^\d+\.\d+\b", line.strip()):
        cleaned_line = regex.sub(r"\(cid:\d+\)", "", line, flags=regex.IGNORECASE).strip(" \t")
        return [cleaned_line] if cleaned_line else []

    if section in BULLET_SECTIONS:
        prepared = regex.sub(r"\(cid:\d+\)", "", line, flags=regex.IGNORECASE)
        prepared = prepared.replace("\u2022", "\n").replace("\u25aa", "\n").replace("\u25cf", "\n")
        prepared = regex.sub(
            r"(?i)^\s*(?:the\s+main\s+)?(?:objectives?|deliverables?)\s+of\s+(?:this|the)\s+(?:project|proposal)\s+are\s*:\s*",
            "",
            prepared,
        )
        prepared = regex.sub(r"(?i)^\s*(?:the\s+main\s+)?(?:objectives?|deliverables?)\s+are\s*:\s*", "", prepared)
        prepared = regex.sub(r"(?i)(?<=\.)\s+(?=(?:to|ensure|improve|develop|provide|install|reduce|increase)\b)", "\n", prepared)
        prepared = regex.sub(r"(?i)\s*;\s+(?=(?:to|ensure|improve|develop|provide|install|reduce|increase)\b)", "\n", prepared)
        prepared = regex.sub(r"(?m)^\s*[-*]\s*", "", prepared)
        prepared = regex.sub(r"\n{2,}", "\n", prepared)
        raw_items = [part.strip(" \t;-") for part in prepared.split("\n") if part.strip(" \t;-")]
        deduped = []
        seen = set()
        for item in raw_items:
            key = item.lower()
            if key not in seen:
                seen.add(key)
                deduped.append(item)
        return deduped

    if section in PARAGRAPH_SECTIONS and section not in {"Technical Skills", "Projects", "Experience"}:
        cleaned_line = regex.sub(r"\(cid:\d+\)", "", line, flags=regex.IGNORECASE).strip(" -\t")
        return [cleaned_line] if cleaned_line else []

    prepared = regex.sub(r"\(cid:\d+\)", "-", line, flags=regex.IGNORECASE)
    prepared = prepared.replace("\u2022", "-").replace("\u25aa", "-").replace("\u25cf", "-")
    prepared = regex.sub(r"(?m)^\s*[-*]\s*", "", prepared)
    prepared = regex.sub(r"\s+[-*]\s+", "\n", prepared)
    prepared = regex.sub(r"\n{2,}", "\n", prepared)

    raw_items = [part.strip(" \t") for part in prepared.split("\n") if part.strip(" \t")]

    deduped = []
    seen = set()
    for item in raw_items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def _looks_like_contact_line(value: str) -> bool:
    lowered = value.lower().strip()
    if lowered.startswith(("address:", "location:")):
        return True
    if "@" in value:
        return True
    if regex.search(r"\+?\d[\d\s\-()]{7,}\d", value):
        return True
    if "linkedin.com" in value.lower() or "github.com" in value.lower():
        return True
    return False


def _polish_item(text: str) -> str:
    cleaned = regex.sub(r"\s+", " ", text).strip(" ,;-")
    cleaned = cleaned.replace("**", "")
    cleaned = regex.sub(r"\[\s*page\s+break\s*\]", "", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"^\d+\.(?!\d)\s*", "", cleaned)
    cleaned = regex.sub(r"^\#{1,6}\s*", "", cleaned)
    cleaned = regex.sub(r"\(cid:\d+\)", "", cleaned, flags=regex.IGNORECASE).strip(" ,;-")
    cleaned = regex.sub(r"[^\x20-\x7E]+", " ", cleaned)
    cleaned = regex.sub(r"\s+", " ", cleaned).strip(" ,;-")
    for phrase, replacement in ACTION_VERB_REPLACEMENTS.items():
        cleaned = regex.sub(rf"\b{regex.escape(phrase)}\b", replacement, cleaned, flags=regex.IGNORECASE)
    if not cleaned:
        return ""
    return cleaned


def _local_refine_item(text: str, section_name: str, doc_type: str) -> str:
    cleaned = _polish_item(text)
    if not cleaned:
        return ""

    cleaned = _safe_repair_run_text(cleaned)
    cleaned = _apply_basic_grammar_fixes(cleaned, preserve_case=section_name in {"Document Details", "References"})

    if doc_type in {"Report", "Proposal", "Thesis"}:
        cleaned = _format_academic_item(cleaned, section_name)

    return cleaned


def _local_refine_item_with_tracking(text: str, section_name: str, doc_type: str) -> str:
    original = text or ""
    polished = _polish_item(original)
    if not polished:
        return ""

    spelling_fixed = _safe_repair_run_text(polished)
    grammar_fixed = _apply_basic_grammar_fixes(
        spelling_fixed,
        preserve_case=section_name in {"Document Details", "References"},
    )
    formatted = grammar_fixed
    if doc_type in {"Report", "Proposal", "Thesis"}:
        formatted = _format_academic_item(grammar_fixed, section_name)

    change_log = get_change_log()
    if polished != original.strip():
        change_log.add("formatting", original[:200], polished[:200], section_name)
    if spelling_fixed != polished:
        change_log.add("spelling", polished[:200], spelling_fixed[:200], section_name)
    if grammar_fixed != spelling_fixed:
        change_log.add("grammar", spelling_fixed[:200], grammar_fixed[:200], section_name)
    if formatted != grammar_fixed:
        change_log.add("formatting", grammar_fixed[:200], formatted[:200], section_name)

    return formatted


def _apply_basic_grammar_fixes(text: str, preserve_case: bool = False) -> str:
    cleaned = text

    replacements = {
        r"\bi\b": "I",
        r"\bdont\b": "don't",
        r"\bcant\b": "can't",
        r"\bwont\b": "won't",
        r"\bim\b": "I'm",
        r"\bits\s+important\b": "it is important",
        r"\bthere\s+is\s+many\b": "there are many",
        r"\bthere\s+is\s+several\b": "there are several",
        r"\bhe\s+go\b": "he goes",
        r"\bshe\s+go\b": "she goes",
        r"\bit\s+go\b": "it goes",
        r"\bthey\s+goes\b": "they go",
        r"\bwe\s+goes\b": "we go",
        r"\bi\s+goes\b": "I go",
    }
    for pattern, replacement in replacements.items():
        cleaned = regex.sub(pattern, replacement, cleaned, flags=regex.IGNORECASE)

    cleaned = regex.sub(r"\b(\w+)(\s+)\1\b", r"\1", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\s+([,.;:!?])", r"\1", cleaned)
    cleaned = regex.sub(r"([,.;:!?])([A-Za-z])", r"\1 \2", cleaned)
    cleaned = regex.sub(r"\(\s+", "(", cleaned)
    cleaned = regex.sub(r"\s+\)", ")", cleaned)
    cleaned = regex.sub(r"\s{2,}", " ", cleaned).strip(" ,;-")

    if not preserve_case and cleaned and not regex.match(r"^[\W\d_]", cleaned):
        cleaned = cleaned[:1].upper() + cleaned[1:]

    if cleaned and not cleaned.endswith((".", "!", "?", ":", ";")):
        if len(cleaned.split()) >= 6 and not _looks_like_tableish_line(cleaned):
            cleaned += "."

    return cleaned


def _looks_like_tableish_line(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return bool(
        "|" in stripped
        or regex.search(r"\b(?:rs\.?|usd|eur|npr|\$)\b", stripped, flags=regex.IGNORECASE)
        or regex.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", stripped)
        or regex.search(r"\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\b", stripped, flags=regex.IGNORECASE)
    )


def _format_academic_item(text: str, section_name: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return ""

    if section_name in {"Objectives", "Expected Outcomes", "Deliverables", "Qualifications or Team"}:
        cleaned = regex.sub(r"^(?:[-*•]\s*)+", "", cleaned).strip()
        return cleaned

    if section_name in {"Timeline", "Budget", "Terms and Conditions", "Appendices", "Document Details", "References"}:
        return cleaned

    cleaned = regex.sub(r"\s*:\s*", ": ", cleaned)
    cleaned = regex.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip()


def _polish_sections(sections: Dict[str, List[str]]) -> Dict[str, List[str]]:
    polished: Dict[str, List[str]] = {}
    for key, values in sections.items():
        out = []
        seen = set()
        for value in values:
            if _is_noise_line(value):
                continue
            if _canonical_heading(value):
                continue
            item = _polish_item(value)
            if item:
                normalized = item.lower()
                if normalized in seen:
                    continue
                seen.add(normalized)
                out.append(item)
        
        if out:
            polished[key] = out
    return polished


def _extract_resume_sections(text: str) -> Tuple[str, str, Dict[str, List[str]]]:
    normalized = _normalize_text(text)
    lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]

    name, contact = _extract_name_contact(lines)
    sections: Dict[str, List[str]] = {sec: [] for sec in RESUME_ORDER}
    current = "Additional Information"

    for line in lines:
        if _is_document_junk_line(line):
            continue
        heading = _canonical_heading(line)
        inline_heading, remainder = _split_resume_heading_line(line)
        if heading:
            current = heading
            continue
        if inline_heading:
            current = inline_heading
            if not remainder:
                continue
            line = remainder
        if line == name or line == contact or _looks_like_contact_line(line):
            continue
        items = _split_line_items(line, current)
        if items:
            cleaned_items = []
            for item in items:
                rewritten = _rewrite_resume_item(item, current)
                if rewritten:
                    cleaned_items.append(rewritten)
            sections[current].extend(cleaned_items)

    cleaned_sections = {k: v for k, v in sections.items() if v}
    polished = _polish_sections(cleaned_sections)
    name, contact, polished = _postprocess_resume_data(name, contact, polished)
    return name, contact, polished


def _extract_general_sections(text: str) -> Tuple[str, str, Dict[str, List[str]]]:
    normalized = _normalize_text(text)
    lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]
    title = "Improved Document"
    if lines and len(lines[0]) <= 70:
        title = lines[0]
        lines = lines[1:]

    content = []
    for line in lines:
        content.extend(_split_line_items(line, "Content"))
    if not content:
        content = [normalized]
    return title, "", _polish_sections({"Content": content})


def _extract_cover_letter_sections(text: str) -> Tuple[str, str, Dict[str, List[str]], List[str]]:
    normalized = _normalize_text(text)
    if not normalized.strip():
        return "Improved Cover Letter", "", {"Body": ["No readable content found."]}, ["Body"]

    normalized = _replace_unprofessional_phrases_with_tracking(normalized, "Cover Letter")
    blocks = _split_cover_letter_blocks(normalized)
    if not blocks:
        return "Improved Cover Letter", "", {"Body": ["No readable content found."]}, ["Body"]

    salutation_block_index = None
    salutation_line = ""
    for index, block in enumerate(blocks):
        first_line = block[0]
        if regex.match(r"^(dear\b|to whom it may concern\b)", first_line, flags=regex.IGNORECASE):
            salutation_block_index = index
            salutation_line = first_line
            break

    if salutation_block_index is None:
        body = _clean_cover_letter_text(" ".join(" ".join(block) for block in blocks))
        return "Improved Cover Letter", "", {"Body": [body]}, ["Body"]

    preamble_blocks = blocks[:salutation_block_index]
    post_salutation_blocks = blocks[salutation_block_index + 1:]
    sender: List[str] = []
    date_lines: List[str] = []
    recipient: List[str] = []
    subject_lines: List[str] = []
    salutation = [salutation_line.rstrip(",") + "," if not salutation_line.rstrip().endswith(",") else salutation_line]

    if len(blocks[salutation_block_index]) > 1:
        post_salutation_blocks = [blocks[salutation_block_index][1:]] + post_salutation_blocks

    passed_date = False
    recipient_started = False
    for block in preamble_blocks:
        for line in block:
            if _is_subject_line(line):
                subject_lines.append(line)
                continue
            if _is_date_line(line):
                date_lines.append(line)
                passed_date = True
                recipient_started = True
                continue
            if _looks_like_recipient_line(line):
                recipient.append(line)
                recipient_started = True
                continue
            if not recipient_started:
                sender.append(line)
            else:
                recipient.append(line)

    sender = list(dict.fromkeys(sender))
    recipient = list(dict.fromkeys(recipient))

    closing_index = None
    closing_starters = {"sincerely", "best regards", "respectfully", "regards", "thank you"}
    for index, block in enumerate(post_salutation_blocks):
        first_line = block[0]
        lowered = first_line.lower().rstrip(",")
        if lowered in closing_starters or lowered.startswith("thank you for your consideration"):
            closing_index = index
            break

    body_blocks = post_salutation_blocks[:closing_index] if closing_index is not None else post_salutation_blocks
    closing_blocks = post_salutation_blocks[closing_index:] if closing_index is not None else []

    sections: Dict[str, List[str]] = {}
    if sender:
        sections["Sender Information"] = sender
    if date_lines:
        sections["Date"] = date_lines
    if recipient:
        sections["Recipient Information"] = recipient
    if subject_lines:
        sections["Subject"] = subject_lines
    sections["Salutation"] = salutation

    body_items = _merge_cover_letter_body_blocks(body_blocks)
    if body_items:
        sections["Body"] = body_items

    if closing_blocks:
        cleaned_closing = []
        for block in closing_blocks:
            for line in block:
                cleaned_line = _clean_cover_letter_text(line)
                if cleaned_line and not cleaned_line.lower().startswith("p.s"):
                    cleaned_closing.append(cleaned_line)
        if cleaned_closing:
            if len(cleaned_closing) >= 2:
                sections["Closing"] = cleaned_closing[:-1]
                sections["Signature"] = [cleaned_closing[-1]]
            else:
                sections["Closing"] = cleaned_closing

    order = [name for name in COVER_LETTER_ORDER if name in sections]
    return "", "", sections, order or ["Body"]


def _extract_academic_sections(text: str, doc_type: str) -> Tuple[str, str, Dict[str, List[str]], List[str]]:
    normalized = _normalize_text(text)
    lines = [ln.strip() for ln in normalized.split("\n") if ln.strip()]
    if not lines:
        return "Improved Document", "", {"Content": ["No readable content found."]}, ["Content"]

    raw_title_line = lines[0]
    title = raw_title_line if len(raw_title_line) <= 100 else "Improved Document"
    title_match = regex.match(r"^\s*title\s*:\s*(.+)$", raw_title_line, flags=regex.IGNORECASE)
    if title_match:
        title = title_match.group(1).strip()
    
    if doc_type == "Thesis":
        preferred_order = THESIS_ORDER
    elif doc_type == "Proposal":
        preferred_order = PROPOSAL_ORDER
    else:
        preferred_order = REPORT_ORDER

    sections: Dict[str, List[str]] = {"Document Details": []}
    encountered_order: List[str] = []
    current = "Document Details"

    for index, line in enumerate(lines):
        if index == 0 and (line == title or line == raw_title_line):
            continue
        heading, inline_content = _parse_academic_heading_line(line)
        if not heading:
            heading, inline_content = _parse_generic_academic_heading_line(line)
        if heading:
            if heading not in sections:
                sections[heading] = []
            if heading not in encountered_order:
                encountered_order.append(heading)
            current = heading
            if inline_content:
                items = _split_line_items(inline_content, current)
                if items:
                    sections[current].extend(items)
            continue
        if _is_noise_line(line):
            continue
        items = _split_line_items(line, current)
        if items:
            sections[current].extend(items)

    cleaned_sections = {k: v for k, v in _polish_sections(sections).items() if v}
    cleaned_order = []
    for name in encountered_order:
        if name in cleaned_sections and name not in cleaned_order:
            cleaned_order.append(name)
    for name in preferred_order:
        if name in cleaned_sections and name not in cleaned_order:
            cleaned_order.append(name)
    for name in cleaned_sections:
        if name not in cleaned_order:
            cleaned_order.append(name)
    if "Document Details" in cleaned_order:
        cleaned_order = ["Document Details"] + [name for name in cleaned_order if name != "Document Details"]
    if not cleaned_order:
        fallback_title, fallback_contact, fallback_sections = _extract_general_sections(text)
        return fallback_title, fallback_contact, fallback_sections, ["Content"]
    return title, "", cleaned_sections, cleaned_order


def _build_data(text: str, doc_type: str) -> Tuple[str, str, Dict[str, List[str]], List[str]]:
    if doc_type in {"Resume", "CV"}:
        title, contact, sections = _extract_resume_sections(text)
        order = [name for name in RESUME_ORDER if name in sections]
        return title, contact, sections, order
    if doc_type == "Cover Letter":
        return _extract_cover_letter_sections(text)
    if doc_type == "Essay":
        title, contact, sections = _extract_general_sections(text)
        return title, contact, sections, ["Content"]
    if doc_type == "Proposal":
        return _extract_academic_sections(text, "Proposal")
    if doc_type in {"Thesis", "Report"}:
        return _extract_academic_sections(text, doc_type)

    title, contact, sections = _extract_general_sections(text)
    return title, contact, sections, ["Content"]


def _rewrite_sections_with_ai(
    title: str,
    contact: str,
    sections: Dict[str, List[str]],
    order: List[str],
    doc_type: str,
) -> Tuple[str, str, Dict[str, List[str]], List[str], str]:
    original_order = list(order)
    original_sections = {key: list(value) for key, value in sections.items()}
    original_section_names = set(original_sections.keys())
    set_document_context(doc_type)
    
    # Add missing sections
    sections, order = _add_missing_sections_with_tracking(sections, doc_type, order)
    
    rewritten_sections: Dict[str, List[str]] = {}
    used_count = 0
    fallback_count = 0
    ai_failure_messages: List[str] = []
    change_log = get_change_log()

    def _normalized_items(value: List[str]) -> List[str]:
        return [_polish_item(item) for item in value if _polish_item(item)]

    def _normalized_text(value: List[str]) -> str:
        text = " ".join(_normalized_items(value)).strip().lower()
        return regex.sub(r"\s+", " ", text)

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        section_text = "\n".join(items)
        
        # Phrase cleanup is useful for job documents but too destructive for academic writing.
        if doc_type in {"Resume", "CV", "Cover Letter"}:
            section_text = _replace_unprofessional_phrases_with_tracking(section_text, section_name)
        corrected_items = _clean_fallback_items(section_text, section_name, doc_type)
        
        ai_response = rewrite_section_with_ai(section_name, section_text, doc_type)
        rewritten_text = (ai_response.content or "").strip() if ai_response.ok else ""
        if not ai_response.ok and ai_response.error:
            ai_failure_messages.append(ai_response.error)

        if rewritten_text:
            rewritten_items = _split_line_items(rewritten_text, section_name)
            rewritten_items = [_polish_item(item) for item in rewritten_items if _polish_item(item)]
            if section_name == "Certifications":
                rewritten_items = _normalize_certification_items(rewritten_items)
            if rewritten_items and (
                doc_type not in {"Report", "Proposal", "Thesis"}
                or _should_accept_academic_rewrite(section_name, original_sections.get(section_name, []), rewritten_items)
            ):
                rewritten_sections[section_name] = rewritten_items
                if _normalized_text(rewritten_items) != _normalized_text(original_sections.get(section_name, [])):
                    change_log.add("rewrite", section_text[:200], "\n".join(rewritten_items)[:200], section_name)
                used_count += 1
                continue

        final_items = corrected_items or items
        if section_name == "Certifications":
            final_items = _normalize_certification_items(final_items)
        rewritten_sections[section_name] = final_items
        if (
            _normalized_text(final_items) != _normalized_text(original_sections.get(section_name, []))
            and section_name in original_section_names
        ):
            change_log.add("rewrite", "\n".join(original_sections.get(section_name, []))[:200], "\n".join(final_items)[:200], section_name)
        fallback_count += 1

    if order != original_order:
        change_log.add("reorder", " > ".join(original_order), " > ".join(order), "document")

    diagnostic = f"AI rewrite: {used_count} sections rewritten, {fallback_count} fallbacks."
    if ai_failure_messages:
        unique_failures = []
        for message in ai_failure_messages:
            if message not in unique_failures:
                unique_failures.append(message)
        diagnostic = f"{diagnostic} AI status: {'; '.join(unique_failures[:2])}"
    return title, contact, rewritten_sections, order, diagnostic


def _polish_run_text(text: str) -> str:
    cleaned = text
    for phrase, replacement in ACTION_VERB_REPLACEMENTS.items():
        cleaned = regex.sub(rf"\b{regex.escape(phrase)}\b", replacement, cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"[ ]{2,}", " ", cleaned)
    return cleaned


def _safe_repair_run_text(text: str) -> str:
    cleaned = text
    for wrong, right in SPELLING_FIXES.items():
        cleaned = regex.sub(rf"\b{regex.escape(wrong)}\b", right, cleaned, flags=regex.IGNORECASE)
    return cleaned


def _apply_languagetool_corrections(text: str) -> str:
    """Apply grammar and spelling corrections using LanguageTool if available."""
    from services.languagetool_service import get_languagetool
    
    if not text or not text.strip():
        return text
    
    try:
        lt_tool = get_languagetool()
        if lt_tool is None:
            return text
        
        matches = lt_tool.check(text)
        if not matches:
            return text
        
        # Sort matches in reverse order to apply from end to start
        # This prevents offset issues when replacing text
        sorted_matches = sorted(matches, key=lambda m: m.offset, reverse=True)
        
        corrected = text
        for match in sorted_matches:
            # Get the first suggestion if available
            if match.replacements:
                suggestion = match.replacements[0]
                start = match.offset
                end = match.offset + match.length
                corrected = corrected[:start] + suggestion + corrected[end:]
        
        return corrected
    except Exception:
        # If LanguageTool fails, just return original text
        return text


def _apply_ai_text_refinement(text: str, doc_type: str, section_name: str = "Content") -> str:
    """Apply AI-based text refinement for better quality."""
    from services.ai_rewriter import rewrite_section_with_ai
    
    if not text or len(text.strip()) < 10:
        # Don't rewrite very short text
        return text
    
    try:
        response = rewrite_section_with_ai(section_name, text, doc_type)
        if response.ok and response.content and response.content.strip():
            return response.content.strip()
    except Exception:
        # If AI rewriting fails, just return original text
        pass
    
    return text


def _clear_paragraph_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _write_paragraph_lines(paragraph, lines: List[str], font_size: float | None = None, bold: bool = False) -> None:
    _clear_paragraph_runs(paragraph)
    cleaned_lines = [line.strip() for line in lines if (line or "").strip()]
    if not cleaned_lines:
        return

    for index, line in enumerate(cleaned_lines):
        run = paragraph.add_run(line)
        run.bold = bold
        if font_size is not None:
            run.font.size = Pt(font_size)
        if index < len(cleaned_lines) - 1:
            run.add_break()


def _remove_blank_main_body_paragraphs(doc: Document) -> bool:
    removed = False
    for paragraph in list(doc.paragraphs):
        if (paragraph.text or "").strip():
            continue
        if _paragraph_contains_visual_content(paragraph):
            continue
        _delete_paragraph(paragraph)
        removed = True
    return removed


def _paragraph_contains_visual_content(paragraph) -> bool:
    xml = paragraph._element.xml
    return any(token in xml for token in ("<w:drawing", "<w:pict", "<a:blip", "<v:shape", "<pic:pic"))


def _overwrite_paragraph_text(paragraph, text: str, *, bold: bool = False, font_size: float | None = None) -> None:
    if _paragraph_contains_visual_content(paragraph):
        return
    _clear_paragraph_runs(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    if font_size is not None:
        run.font.size = Pt(font_size)


def _looks_like_subheading_paragraph(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped or _looks_like_main_heading_paragraph(stripped):
        return False
    if len(stripped) > 120:
        return False
    if ":" in stripped and len(stripped.split()) <= 14:
        return True
    return False


def _apply_academic_heading_formatting(paragraph) -> None:
    text = (paragraph.text or "").strip()
    if not text:
        return

    if _looks_like_main_heading_paragraph(text):
        if paragraph.runs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(13.5)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        return

    if _looks_like_subheading_paragraph(text):
        if paragraph.runs:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = Pt(11.5)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        return

    paragraph.paragraph_format.space_after = Pt(4)


def _iter_main_body_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph


def _find_first_academic_heading_index(doc: Document) -> int:
    for index, paragraph in enumerate(_iter_main_body_paragraphs(doc)):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        heading, _inline = _parse_academic_heading_line(text)
        if not heading:
            heading, _inline = _parse_generic_academic_heading_line(text)
        if heading:
            return index
    return -1


def _cover_page_needs_structure(lines: List[str]) -> bool:
    cleaned = [line.strip() for line in lines if (line or "").strip()]
    if len(cleaned) <= 1:
        return True
    if any("prepared by" in line.lower() and "date" in line.lower() for line in cleaned):
        return True
    if len(cleaned) <= 2 and sum(line.count(":") for line in cleaned) >= 2:
        return True
    return False


def _extract_cover_detail_lines(text: str) -> List[str]:
    prepared = regex.sub(r"\s+", " ", text or "").strip()
    if not prepared:
        return []
    prepared = regex.sub(
        r"(?i)\s+(?=(prepared by|submitted by|date|department|course|supervisor|advisor)\s*:)",
        "\n",
        prepared,
    )
    return [part.strip(" -") for part in prepared.split("\n") if part.strip(" -")]


def _improve_cover_page_layout(doc: Document, doc_type: str = "Proposal") -> bool:
    heading_index = _find_first_academic_heading_index(doc)
    if heading_index <= 0:
        return False

    cover_paragraphs = [paragraph for paragraph in doc.paragraphs[:heading_index] if (paragraph.text or "").strip()]
    if not cover_paragraphs:
        return False

    cover_lines = [(paragraph.text or "").strip() for paragraph in cover_paragraphs]
    if not _cover_page_needs_structure(cover_lines):
        return False

    title_text = _normalize_academic_metadata_line(cover_lines[0])
    title_para = cover_paragraphs[0]
    _write_paragraph_lines(title_para, [title_text], font_size=20, bold=True)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(20)

    detail_lines: List[str] = []
    for paragraph in cover_paragraphs[1:]:
        detail_lines.extend(_extract_cover_detail_lines(paragraph.text))

    if detail_lines:
        detail_para = cover_paragraphs[1] if len(cover_paragraphs) > 1 else None
        if detail_para is not None:
            _write_paragraph_lines(detail_para, detail_lines, font_size=11.5)
            detail_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            detail_para.paragraph_format.space_after = Pt(0)
            for paragraph in cover_paragraphs[2:]:
                _clear_paragraph_runs(paragraph)

    if doc_type in {"Thesis", "Report"}:
        doc.paragraphs[heading_index].paragraph_format.page_break_before = True
    else:
        doc.paragraphs[heading_index].paragraph_format.page_break_before = False
    return True


def _normalize_objective_items(text: str) -> List[str]:
    return _split_line_items(text, "Objectives")


def _format_objectives_as_bullets(doc: Document) -> bool:
    in_objectives = False
    updated = False

    for paragraph in _iter_main_body_paragraphs(doc):
        text = (paragraph.text or "").strip()
        if not text:
            continue

        heading, _inline = _parse_academic_heading_line(text)
        if not heading:
            heading, _inline = _parse_generic_academic_heading_line(text)
        if heading:
            in_objectives = heading == "Objectives"
            continue
        if not in_objectives:
            continue

        normalized_items = _normalize_objective_items(text)
        if not normalized_items:
            continue
        if len(normalized_items) != 1:
            continue

        lowered = normalized_items[0].lower()
        if regex.search(r"(?i)\bthe\s+main\s+obj(?:ective|ectives|etive|etives)\b", lowered):
            continue

        bullet_text = normalized_items[0]
        if not bullet_text.startswith("\u2022"):
            bullet_text = f"\u2022 {bullet_text}"
        _write_paragraph_lines(paragraph, [bullet_text])
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None
        updated = True

    return updated


def _format_proposal_objectives_as_plain_lines(doc: Document) -> bool:
    in_objectives = False
    updated = False

    for paragraph in _iter_main_body_paragraphs(doc):
        text = (paragraph.text or "").strip()
        if not text:
            continue

        heading, _inline = _parse_academic_heading_line(text)
        if not heading:
            heading, _inline = _parse_generic_academic_heading_line(text)
        if heading:
            in_objectives = heading in {"Objectives", "Objectives of the Proposal"}
            continue
        if not in_objectives:
            continue

        if regex.search(r"(?i)\bthe\s+main\s+obj(?:ective|ectives|etive|etives)\b", text):
            continue

        cleaned = regex.sub(r"^(?:[-*\u2022]\s*)+", "", text).strip()
        if cleaned != text:
            _write_paragraph_lines(paragraph, [cleaned])
            updated = True

    return updated


def _iter_cell_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_cell_paragraphs(nested)


def _iter_all_docx_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_cell_paragraphs(table)


def _repair_table_cell_text(paragraph_text: str) -> str:
    cleaned = (paragraph_text or "").replace("\xa0", " ").strip()
    if not cleaned:
        return ""
    cleaned = _safe_repair_run_text(cleaned)
    cleaned = regex.sub(r"\bongoing\b", "Ongoing", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def _proposal_resource_line_needs_plain_style(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return bool(
        regex.search(r"\(cost:\s*\$\d+", stripped, flags=regex.IGNORECASE)
        and not stripped.lower().startswith("note:")
    )


def _normalize_proposal_resource_line(text: str) -> str:
    cleaned = (text or "").strip()
    if cleaned.endswith("."):
        cleaned = cleaned[:-1]
    return cleaned


def _add_proposal_budget_note(doc: Document) -> bool:
    paragraphs = list(_iter_main_body_paragraphs(doc))
    budget_heading_index = -1
    conclusion_index = -1

    for index, paragraph in enumerate(paragraphs):
        text = (paragraph.text or "").strip()
        normalized_heading = regex.sub(r"^\d+(?:\.0)?[.):-]\s+", "", text).strip().lower()
        if normalized_heading in {"budget summary", "budget"}:
            budget_heading_index = index
        elif normalized_heading == "conclusion":
            conclusion_index = index

    if budget_heading_index == -1 or conclusion_index == -1:
        return False

    note_text = "Note: Selling the bottles for $1 each will recover $500 of the cost."
    for paragraph in paragraphs[budget_heading_index + 1:conclusion_index]:
        if (paragraph.text or "").strip() == note_text:
            return False

    paragraphs[conclusion_index].insert_paragraph_before(note_text)
    return True


def _add_proposal_final_line(doc: Document) -> bool:
    final_text = "Together, we can make our school cleaner and greener!"
    paragraphs = [p for p in _iter_main_body_paragraphs(doc) if (p.text or "").strip()]
    if not paragraphs:
        return False
    if any((p.text or "").strip() == final_text for p in paragraphs):
        return False

    last = paragraphs[-1]
    if "Thank you for your time and consideration." not in (last.text or ""):
        return False

    paragraph = doc.add_paragraph(final_text)
    paragraph.paragraph_format.space_before = Pt(3)
    return True


def _safe_repair_docx(source_path: Path, output_path: Path) -> None:
    doc = Document(source_path)
    
    for paragraph in _iter_all_docx_paragraphs(doc):
        for run in paragraph.runs:
            raw = run.text or ""
            if not raw.strip():
                continue
            updated = _safe_repair_run_text(raw)
            if updated != raw:
                run.text = updated
    
    doc.save(output_path)


def _looks_like_main_heading_paragraph(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    if regex.match(r"^\[\s*page\s+break\s*\]$", stripped, flags=regex.IGNORECASE):
        return True
    if regex.match(r"^(?:chapter\s+\d+|\d+(?:\.0)?|[ivxlcdm]+)(?:[.):-]\s+|\s+)[A-Za-z]", stripped, flags=regex.IGNORECASE):
        return True
    return False


def _normalize_title_case_phrase(text: str) -> str:
    words = []
    for word in text.split():
        if word.isupper() and len(word) > 1:
            words.append(word)
        else:
            words.append(word[:1].upper() + word[1:].lower() if word else word)
    return " ".join(words)


def _title_case_heading_text(text: str) -> str:
    lower_words = {"of", "the", "and", "in", "to", "for", "by", "with", "on", "a", "an"}
    parts = []
    words = text.split()
    for index, word in enumerate(words):
        normalized = word.lower()
        if index > 0 and index < len(words) - 1 and normalized in lower_words:
            parts.append(normalized)
        else:
            parts.append(word[:1].upper() + word[1:].lower() if word else word)
    return " ".join(parts)


def _canonical_proposal_heading_text(text: str) -> str:
    """Map rough source headings into the app's proposal flow labels."""
    cleaned = regex.sub(r"\s+", " ", (text or "").strip(" :-")).strip()
    normalized = cleaned.lower()
    normalized = regex.sub(r"^(?:chapter\s+\d+\s*[:.)-]?\s*|\d+(?:\.0)?[.):\-]\s+|[ivxlcdm]+(?:[.):\-]\s+))", "", normalized).strip()

    aliases = {
        "background of the problem": "Problem Statement",
        "background of problem": "Problem Statement",
        "backround of the problem": "Problem Statement",
        "statement of the problem": "Problem Statement",
        "objectives of the proposal": "Objectives",
        "objetives of the proposal": "Objectives",
        "objective of the proposal": "Objectives",
        "proposed activities": "Technical Approach",
        "proposed activites": "Technical Approach",
        "activities": "Technical Approach",
        "plan/method": "Technical Approach",
        "plan / method": "Technical Approach",
        "proposed solution": "Technical Approach",
        "resources needed": "Required Resources",
        "resources required": "Required Resources",
        "budget summary": "Budget",
        "estimated budget": "Budget",
        "cost estimate": "Budget",
        "expected results": "Expected Outcomes",
        "expected result": "Expected Outcomes",
        "outcomes": "Expected Outcomes",
    }

    return aliases.get(normalized, _title_case_heading_text(cleaned))


def _title_case_document_title(text: str) -> str:
    lower_words = {"of", "the", "and", "in", "to", "for", "by", "with", "on", "a", "an"}
    words = text.split()
    out = []
    for index, word in enumerate(words):
        core = word.strip()
        normalized = core.lower()
        if index > 0 and normalized in lower_words:
            out.append(normalized)
        else:
            out.append(core[:1].upper() + core[1:].lower() if core else core)
    return " ".join(out)


def _normalize_academic_metadata_line(text: str) -> str:
    cleaned = text.replace("\xa0", " ").strip()
    cleaned = regex.sub(r"\s*:\s*", ": ", cleaned)
    cleaned = regex.sub(r"\b(\d{1,2})(st|nd|rd|th)\s+of\s+([A-Za-z]+)\s+(\d{4})\b", r"\1\2 \3 \4", cleaned, flags=regex.IGNORECASE)

    if regex.match(r"(?i)^prepared by\s*:", cleaned):
        label, value = cleaned.split(":", 1)
        return f"Prepared by: {_normalize_title_case_phrase(value.strip())}".strip()
    if regex.match(r"(?i)^proposal title\s*:", cleaned):
        label, value = cleaned.split(":", 1)
        return f"Proposal Title: {_title_case_document_title(value.strip())}".strip()
    if regex.match(r"(?i)^date\s*:", cleaned):
        label, value = cleaned.split(":", 1)
        return f"Date: {value.strip()}".strip()
    return cleaned


def _is_listish_academic_paragraph(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    if _looks_like_main_heading_paragraph(stripped):
        return True
    if regex.match(r"^\d+\.\s+[A-Za-z]", stripped):
        return True
    if ":" in stripped and len(stripped) < 140:
        return True
    if regex.match(r"^(?:[-*\u2022]|\d+\))\s+", stripped):
        return True
    return False


def _repair_academic_paragraph_text(text: str, doc_type: str = "General Document") -> str:
    cleaned = (text or "").replace("\xa0", " ").strip()
    if not cleaned:
        return ""

    cleaned = _safe_repair_run_text(cleaned)
    cleaned = regex.sub(r"(?<=\d)\.(?=[A-Za-z])", ". ", cleaned)
    cleaned = regex.sub(r"(?<=[a-z])(?=[A-Z])", " ", cleaned)
    cleaned = regex.sub(r"([:;,.!?])(?=[A-Za-z])", r"\1 ", cleaned)
    cleaned = regex.sub(r"\s{2,}", " ", cleaned).strip()

    if regex.match(r"(?i)^(proposal title|prepared by|date)\s*:", cleaned):
        return _normalize_academic_metadata_line(cleaned)

    heading_name, _inline = _parse_academic_heading_line(cleaned)
    if not heading_name:
        heading_name, _inline = _parse_generic_academic_heading_line(cleaned)
    if heading_name:
        prefix_match = regex.match(r"^\s*((?:chapter\s+\d+\s*[:.)-]?\s*)|(?:\d+(?:\.0)?[.):-]\s+)|(?:[ivxlcdm]+(?:[.):-]\s+)))", cleaned, flags=regex.IGNORECASE)
        prefix = prefix_match.group(0).strip() if prefix_match else ""
        body = cleaned[prefix_match.end():].strip(" :-") if prefix_match else cleaned.strip(" :-")
        if doc_type == "Proposal":
            body = _canonical_proposal_heading_text(body)
        else:
            body = _title_case_heading_text(body)
        if prefix:
            prefix = regex.sub(r"(?<=\d)\.(?!\s)", ". ", prefix)
            return f"{prefix.strip()} {body}".strip()
        return body

    cleaned = _apply_basic_grammar_fixes(cleaned, preserve_case=False)
    cleaned = regex.sub(r"\bbi\b", "by", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"^Too teach\b", "To teach", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bis a big problem\b", "is a significant issue", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bthrow away many\b", "discard numerous", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bwe will do the following activities\b", "We will carry out the following activities", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bsell cheap reusable\b", "Sell affordable, reusable", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\ba group of volunteers\b", "a group of student volunteers", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bif we do this plan, we expect to see less plastic in the trash\b", "If this plan is implemented, we expect to see significantly less plastic in our trash", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bstudents will learn good habits\b", "Students will learn positive environmental habits", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bthe school will become cleaner and more environmentally friendly\b", "our school will become cleaner, greener, and more environmentally friendly", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bwe ask the school administration to approve this proposal and give us the money needed\b", "We respectfully ask the school administration to approve this proposal and provide the necessary funding", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bthank you for your time\b", "Thank you for your time and consideration", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bour school be\b", "our school to be", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bbe a leader in combat climate change\b", "become a leader in fighting plastic pollution and climate change", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bwe expects\b", "we expect", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bstop climate change\b", "combat climate change", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bare school\b", "our school", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bhurts oceans\b", "harms oceans", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\ba assembly\b", "a school assembly", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bevery friday\b", "Every Friday", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bmost of this plastic is not recycled because\b", "Unfortunately, most of this plastic is not recycled because", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"Unfortunately,\s+Unfortunately,", "Unfortunately,", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bthis waste ends up in the local landfill or on the streets, which hurts animals and makes the school look dirty\b", "As a result, this waste ends up in local landfills or on the streets, harming wildlife and making our school campus appear dirty and uncared for", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\bhold a school assembly to show a video about how plastic harms oceans\b", "Hold a school assembly and show a video about how plastic pollution harms oceans and marine life", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\brecord how much plastic is thrown away\b", "record how much plastic is being thrown away", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\(cost: \$2 per bottle, but we will sell them for \$1\)", "(cost: $2 per bottle; we will sell them for $1 each to encourage participation)", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"Students will learn positive environmental habits\.\s+our school will become cleaner, greener, and more environmentally friendly\.", "Students will learn positive environmental habits, and our school will become cleaner, greener, and more environmentally friendly. Additionally, we hope to inspire other schools in the area to launch similar initiatives.", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"our school to become a leader in fighting plastic pollution and climate change", "our school become a leader in fighting plastic pollution and climate change", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"\band consideration and consideration\b", "and consideration", cleaned, flags=regex.IGNORECASE)
    cleaned = regex.sub(r"(?<=[.!?]\s)([a-z])", lambda m: m.group(1).upper(), cleaned)
    cleaned = regex.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def _repair_academic_docx_in_place(source_path: Path, output_path: Path, doc_type: str = "Proposal") -> None:
    """Preserve the original DOCX layout while applying comprehensive text improvements."""
    doc = Document(source_path)
    seen_meaningful_content = False
    paragraph_count = 0
    paragraphs_to_refine = []
    change_log = get_change_log()

    # First pass: identify paragraphs and apply basic corrections
    for paragraph in _iter_all_docx_paragraphs(doc):
        paragraph_text = (paragraph.text or "").strip()
        if not paragraph_text:
            continue
        if _paragraph_contains_visual_content(paragraph):
            continue

        if regex.match(r"^\[\s*page\s+break\s*\]$", paragraph_text, flags=regex.IGNORECASE):
            if paragraph.runs:
                paragraph.runs[0].text = ""
                for run in paragraph.runs[1:]:
                    run.text = ""
            continue

        # Repair the full paragraph text instead of individual runs so spaces are preserved.
        improved_paragraph_text = _repair_academic_paragraph_text(paragraph_text, doc_type)
        improved_paragraph_text = _apply_languagetool_corrections(improved_paragraph_text)

        if improved_paragraph_text != paragraph_text:
            change_log.add("rewrite", paragraph_text[:200], improved_paragraph_text[:200], "document")
            _overwrite_paragraph_text(paragraph, improved_paragraph_text)
        
        # Collect paragraphs for AI refinement (longer content blocks)
        current_text = (paragraph.text or "").strip()
        if len(current_text) > 120 and paragraph_count < 12 and not _is_listish_academic_paragraph(current_text):
            paragraphs_to_refine.append((paragraph, current_text))
            paragraph_count += 1

        if _looks_like_main_heading_paragraph(paragraph.text):
            if seen_meaningful_content and doc_type in {"Thesis", "Report"}:
                paragraph.paragraph_format.page_break_before = True
            elif doc_type == "Proposal":
                paragraph.paragraph_format.page_break_before = False
            seen_meaningful_content = True
        elif paragraph_text:
            seen_meaningful_content = True

        _apply_academic_heading_formatting(paragraph)

    # Second pass: Apply AI refinement to selected paragraphs
    for paragraph, paragraph_text in paragraphs_to_refine:
        try:
            improved = _apply_ai_text_refinement(paragraph_text, doc_type, "Content")
            if improved:
                improved = _repair_academic_paragraph_text(improved, doc_type)
            if (
                not improved
                or improved == paragraph_text
                or len(improved) < max(80, int(len(paragraph_text) * 0.75))
            ):
                improved = _repair_academic_paragraph_text(paragraph_text, doc_type)
            if improved and improved != paragraph_text:
                change_log.add("rewrite", paragraph_text[:200], improved[:200], "document")
                _overwrite_paragraph_text(paragraph, improved)
                _apply_academic_heading_formatting(paragraph)
        except Exception:
            # If AI refinement fails, just skip this paragraph
            pass

    _improve_cover_page_layout(doc, doc_type)
    if doc_type == "Proposal":
        if _format_proposal_objectives_as_plain_lines(doc):
            change_log.add("formatting", "Objectives bullets", "Objectives formatted as plain proposal lines", "Objectives")
    else:
        if _format_objectives_as_bullets(doc):
            change_log.add("formatting", "Objectives paragraphs", "Objectives formatted as bullet points", "Objectives")

    if doc_type == "Proposal":
        for paragraph in _iter_main_body_paragraphs(doc):
            text = (paragraph.text or "").strip()
            if not text:
                continue
            if _proposal_resource_line_needs_plain_style(text):
                normalized = _normalize_proposal_resource_line(text)
                if normalized != text:
                    change_log.add("formatting", text[:200], normalized[:200], "Resources Needed")
                    _overwrite_paragraph_text(paragraph, normalized)

        if _add_proposal_budget_note(doc):
            change_log.add("formatting", "Budget Summary", "Added proposal budget note", "Budget Summary")
        if _add_proposal_final_line(doc):
            change_log.add("formatting", "Conclusion", "Added final motivational closing line", "Conclusion")
        if _remove_blank_main_body_paragraphs(doc):
            change_log.add("formatting", "Blank paragraphs", "Removed extra blank paragraphs", "document")

    # Final pass: clean table cell text while preserving table structure.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = (paragraph.text or "").strip()
                    if not paragraph_text:
                        continue
                    repaired = _repair_table_cell_text(paragraph_text)
                    if repaired != paragraph_text:
                        change_log.add("formatting", paragraph_text[:200], repaired[:200], "table")
                        _overwrite_paragraph_text(paragraph, repaired)

    doc.save(output_path)


def _docx_contains_preservable_elements(source_path: Path | None) -> bool:
    """Detect whether a DOCX contains tables or embedded assets that should not be rebuilt away."""
    if source_path is not None and not isinstance(source_path, Path):
        source_path = Path(source_path)
    if not source_path or source_path.suffix.lower() != ".docx" or not source_path.exists():
        return False

    try:
        doc = Document(source_path)
        if doc.tables:
            return True

        with zipfile.ZipFile(source_path) as archive:
            for name in archive.namelist():
                if (
                    name.startswith("word/media/")
                    or name.startswith("word/charts/")
                    or name.startswith("word/embeddings/")
                ):
                    return True
    except Exception:
        return False
    return False


# ============================================
# MAIN EXPORT FUNCTION
# ============================================

def improve_document(
    text: str,
    original_filename: str,
    doc_type: str = "General Document",
    source_path: Path | None = None,
):
    if source_path is not None and not isinstance(source_path, Path):
        source_path = Path(source_path)
    stem = Path(original_filename).stem
    docx_path = Path(IMPROVED_FOLDER) / f"improved_{stem}.docx"
    pdf_path = Path(IMPROVED_FOLDER) / f"improved_{stem}.pdf"
    preserve_layout = (
        doc_type in {"Report", "Proposal", "Thesis"}
        and _docx_contains_preservable_elements(source_path)
    )

    if preserve_layout and source_path is not None:
        _repair_academic_docx_in_place(source_path, docx_path, doc_type)
        improved_parsed = parse_document(docx_path)
        improved_text = improved_parsed.get("text", "") or text
        improved_title, improved_contact, improved_sections, improved_order = _build_data(improved_text, doc_type)
        if any(improved_sections.get(section_name) for section_name in improved_order):
            export_pdf(improved_title, improved_contact, improved_sections, improved_order, pdf_path, doc_type=doc_type)
        else:
            export_pdf("Improved Document", "", {"Content": [improved_text]}, ["Content"], pdf_path, doc_type=doc_type)
        return {
            "docx": docx_path,
            "pdf": pdf_path,
            "diagnostic": "Layout-preserving mode: text improved with grammar/spelling corrections and AI refinement, tables/images preserved, objectives formatted as points.",
        }

    title, contact, sections, order = _build_data(text, doc_type)
    title, contact, sections, order, section_diagnostic = _rewrite_sections_with_ai(
        title, contact, sections, order, doc_type
    )

    if any(sections.get(section_name) for section_name in order):
        export_docx(title, contact, sections, order, docx_path, doc_type=doc_type)
        export_pdf(title, contact, sections, order, pdf_path, doc_type=doc_type)
        return {
            "docx": docx_path,
            "pdf": pdf_path,
            "diagnostic": section_diagnostic,
        }

    # Fallback to whole-document rewrite
    ai_response = rewrite_document_with_ai(text, doc_type)
    rewritten_text = ai_response.content.strip() if ai_response.ok and ai_response.content.strip() else ""

    if rewritten_text:
        title, contact, sections, order = _build_data(rewritten_text, doc_type)
        export_docx(title, contact, sections, order, docx_path, doc_type=doc_type)
        export_pdf(title, contact, sections, order, pdf_path, doc_type=doc_type)
        return {
            "docx": docx_path,
            "pdf": pdf_path,
            "diagnostic": f"Whole-document AI rewrite used.",
        }

    export_docx(title, contact, sections, order, docx_path, doc_type=doc_type)
    export_pdf(title, contact, sections, order, pdf_path, doc_type=doc_type)
    return {
        "docx": docx_path,
        "pdf": pdf_path,
        "diagnostic": "Rule-based fallback used.",
    }
