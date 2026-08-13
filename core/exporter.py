"""
Document Export Service
Handles DOCX and PDF export for all document types with proper formatting
"""

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer
import regex

from config.format_profiles import get_format_profile

# ============================================
# CONSTANTS
# ============================================

PARAGRAPH_SECTIONS = {
    "Document Details",
    "Acknowledgement",
    "Abstract",
    "Executive Summary",
    "Introduction",
    "Methodology",
    "Analysis",
    "Results",
    "Discussion",
    "Findings",
    "Conclusion",
    "Recommendations",
    "References",
    "Professional Summary",
    "Content",
}

RESUME_STRUCTURED_SECTIONS = {"Experience", "Projects", "Education"}

COMMON_RESUME_ACTION_VERBS = {
    "achieved", "administered", "analyzed", "built", "collaborated", "configured",
    "created", "delivered", "designed", "developed", "drove", "enhanced", "executed",
    "implemented", "improved", "increased", "launched", "led", "maintained",
    "managed", "optimized", "organized", "performed", "prepared", "produced",
    "reduced", "resolved", "streamlined", "supported", "trained",
}


def _split_resume_contact_lines(contact: str) -> List[str]:
    parts = [part.strip() for part in (contact or "").split("|") if part.strip()]
    if len(parts) <= 1:
        return parts
    if len(contact) <= 42:
        return [" | ".join(parts)]
    return parts


def _education_item_is_institution(item: str) -> bool:
    lowered = (item or "").lower()
    institution_keywords = {"university", "college", "school", "institute", "academy", "campus"}
    return any(keyword in lowered for keyword in institution_keywords)


def _experience_item_should_be_header(item: str) -> bool:
    text = (item or "").strip()
    if not text:
        return False
    if text.endswith((".", "!", "?")):
        return False
    words = text.split()
    if len(words) > 8:
        return False
    title_case_ratio = (
        sum(1 for word in words if regex.fullmatch(r"[A-Z][A-Za-z'`.-]*", word.strip(",.:;()")))
        / max(1, len(words))
    )
    has_date = bool(
        regex.search(
            r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\b|\b(?:19|20)\d{2}\b|\bpresent\b|\bcurrent\b",
            text.lower(),
            flags=regex.IGNORECASE,
        )
    )
    has_location = "," in text
    return title_case_ratio >= 0.55 or has_date or has_location


# ============================================
# HELPER FUNCTIONS
# ============================================

def _escape_html(text: str) -> str:
    """Escape HTML special characters for PDF generation"""
    if not text:
        return ""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _shade_paragraph(paragraph, fill: str = "EAF1FF") -> None:
    """Add background shading to a paragraph"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _add_horizontal_rule_docx(doc: Document, color: str = "D7DFED") -> None:
    """Add a light horizontal divider line as its own paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    p_pr.append(border)


def _add_cover_page_docx(doc: Document, title: str, details: List[str], accent: RGBColor) -> bool:
    """Render an existing academic cover page before the main document body."""
    cover_lines = [line.strip() for line in details if (line or "").strip()]
    if not title and not cover_lines:
        return False

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(140)

    if title:
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(18)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = accent

    for index, line in enumerate(cover_lines):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10 if index < len(cover_lines) - 1 else 0)
        run = paragraph.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()
    return True


def _resume_line_type(section_name: str, item: str) -> str:
    """Classify resume lines so structured entries are not flattened into bullets."""
    text = (item or "").strip()
    if not text:
        return "empty"

    words = text.split()
    word_count = len(words)
    lowered = text.lower()
    first_word = words[0].lower().strip(",.:;()") if words else ""

    has_date = bool(
        regex.search(
            r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\b|\b(?:19|20)\d{2}\b|\bpresent\b|\bcurrent\b",
            lowered,
            flags=regex.IGNORECASE,
        )
    )
    has_meta_separator = any(token in text for token in {"|", "@", " - ", " – ", " — "})
    ends_like_sentence = text.endswith((".", "!", "?"))

    if section_name == "Education":
        education_keywords = {
            "university", "college", "school", "institute", "academy",
            "bachelor", "master", "phd", "gpa", "certificate", "diploma",
        }
        if any(keyword in lowered for keyword in education_keywords):
            return "detail" if ends_like_sentence and word_count > 10 else "header"

    if section_name in RESUME_STRUCTURED_SECTIONS:
        if section_name == "Experience" and not has_meta_separator and not has_date and word_count >= 6:
            return "bullet"
        if first_word in COMMON_RESUME_ACTION_VERBS:
            return "bullet"
        if ends_like_sentence and word_count >= 4:
            return "bullet"
        if has_meta_separator or has_date:
            return "header"
        if word_count <= 7:
            return "header"
        return "detail"

    return "bullet"


def _add_resume_entry_line_docx(doc: Document, item: str, line_type: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2.5)

    if line_type == "header":
        run = paragraph.add_run(item)
        run.bold = True
        run.font.size = Pt(10.8)
        return

    if line_type == "detail":
        paragraph.paragraph_format.left_indent = Inches(0.15)
        run = paragraph.add_run(item)
        run.font.size = Pt(10.4)
        return

    bullet_run = paragraph.add_run("\u2022 ")
    bullet_run.font.color.rgb = RGBColor(34, 72, 127)
    text_run = paragraph.add_run(item)
    text_run.font.size = Pt(10.5)


def _add_resume_entry_line_pdf(story: list, item: str, line_type: str, body_style: ParagraphStyle) -> None:
    if line_type == "header":
        header_style = ParagraphStyle(
            "resume_entry_header",
            parent=body_style,
            fontName="Helvetica-Bold",
            fontSize=max(body_style.fontSize, 10.3),
            spaceAfter=2,
        )
        story.append(Paragraph(_escape_html(item), header_style))
        return

    if line_type == "detail":
        detail_style = ParagraphStyle(
            "resume_entry_detail",
            parent=body_style,
            leftIndent=10,
            spaceAfter=2,
        )
        story.append(Paragraph(_escape_html(item), detail_style))
        return

    bullet_style = ParagraphStyle(
        "resume_entry_bullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=-10,
        bulletIndent=0,
        spaceAfter=2.5,
    )
    story.append(Paragraph(_escape_html(item), bullet_style, bulletText="\u2022"))


def _setup_docx(doc: Document, profile: Dict) -> None:
    """Configure DOCX document margins and default font"""
    section = doc.sections[0]
    margins = profile.get("docx_margins", {})
    section.top_margin = Inches(margins.get("top", 0.8))
    section.bottom_margin = Inches(margins.get("bottom", 0.8))
    section.left_margin = Inches(margins.get("left", 0.85))
    section.right_margin = Inches(margins.get("right", 0.85))

    docx_profile = profile.get("docx", {})
    normal = doc.styles["Normal"]
    normal.font.name = docx_profile.get("normal_font", "Calibri")
    normal.font.size = Pt(docx_profile.get("normal_size", 11))


def _needs_page_break(section_name: str, page_break_before: set, first_section_written: bool) -> bool:
    """Determine if a page break is needed before a section"""
    return first_section_written and section_name in page_break_before


def _add_general_header_docx(doc: Document, title: str, contact: str) -> None:
    """Add simple header for general documents"""
    if title and title != "Improved Document":
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(6)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
    
    if contact:
        contact_paragraph = doc.add_paragraph(contact)
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.paragraph_format.space_after = Pt(8)


# ============================================
# COVER LETTER EXPORT
# ============================================

def _add_cover_letter_docx(
    doc: Document,
    title: str,
    contact: str,
    sections: Dict[str, List[str]],
    order: List[str],
    body_space_after: float,
) -> None:
    """Add cover letter content to DOCX"""
    def add_block(lines: List[str], space_after: float = 0, bold_last: bool = False) -> None:
        last_index = len(lines) - 1
        for index, line in enumerate(lines):
            paragraph = doc.add_paragraph(line)
            paragraph.paragraph_format.space_after = Pt(space_after)
            if bold_last and index == last_index and paragraph.runs:
                paragraph.runs[0].bold = True

    if title and title != "Improved Document":
        add_block([title], space_after=0)
    if contact:
        add_block([part.strip() for part in contact.split("|") if part.strip()], space_after=0)
    if title or contact:
        doc.add_paragraph()

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        if section_name in {"Sender Information", "Date", "Recipient Information"}:
            add_block(items, space_after=0)
            doc.add_paragraph()
            continue

        if section_name == "Subject":
            for item in items:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(10)
                run = paragraph.add_run(item)
                run.bold = True
            continue

        if section_name == "Salutation":
            add_block(items, space_after=10)
            continue

        if section_name == "Body":
            for item in items:
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.space_after = Pt(max(body_space_after, 10))
            continue

        if section_name == "Closing":
            closing_lines = list(items)
            closing_lines = ["Sincerely," if line.lower() == "sincerely" else line for line in closing_lines]
            add_block(closing_lines, space_after=0)
            doc.add_paragraph()
            continue

        if section_name == "Signature":
            add_block(items, space_after=0, bold_last=True)
            doc.add_paragraph()
            continue

        add_block(items, space_after=0)
        doc.add_paragraph()


def _add_cover_letter_pdf(
    story: list,
    title: str,
    contact: str,
    sections: Dict[str, List[str]],
    order: List[str],
    body_style: ParagraphStyle,
) -> None:
    """Add cover letter content to PDF"""
    meta_style = ParagraphStyle(
        "cover_letter_meta",
        parent=body_style,
        spaceAfter=2,
    )
    subject_style = ParagraphStyle(
        "cover_letter_subject",
        parent=body_style,
        fontName="Helvetica-Bold",
        spaceAfter=10,
    )
    signoff_style = ParagraphStyle(
        "cover_letter_signoff",
        parent=body_style,
        spaceAfter=2,
    )
    signature_style = ParagraphStyle(
        "cover_letter_signature",
        parent=body_style,
        fontName="Helvetica-Bold",
        spaceAfter=8,
    )

    def add_lines(lines: List[str], gap_after: float = 0.1, style: ParagraphStyle | None = None) -> None:
        paragraph_style = style or body_style
        for line in lines:
            story.append(Paragraph(_escape_html(line), paragraph_style))
        story.append(Spacer(1, gap_after * inch))

    if title and title != "Improved Document":
        add_lines([title], gap_after=0)
    if contact:
        add_lines([part.strip() for part in contact.split("|") if part.strip()], gap_after=0)
    if title or contact:
        story.append(Spacer(1, 0.12 * inch))

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        if section_name in {"Sender Information", "Date", "Recipient Information"}:
            add_lines(items, gap_after=0.12, style=meta_style)
            continue

        if section_name == "Subject":
            add_lines(items, gap_after=0.12, style=subject_style)
            continue

        if section_name == "Salutation":
            add_lines(items, gap_after=0.1, style=body_style)
            continue

        if section_name == "Body":
            for item in items:
                story.append(Paragraph(_escape_html(item), body_style))
                story.append(Spacer(1, 0.12 * inch))
            continue

        if section_name == "Closing":
            closing_lines = list(items)
            closing_lines = ["Sincerely," if line.lower() == "sincerely" else line for line in closing_lines]
            if closing_lines:
                story.append(Paragraph(_escape_html(closing_lines[0]), signoff_style))
            story.append(Spacer(1, 0.12 * inch))
            continue

        if section_name == "Signature":
            for line in items:
                story.append(Paragraph(_escape_html(line), signature_style))
            story.append(Spacer(1, 0.12 * inch))
            continue

        add_lines(items, gap_after=0.12)


# ============================================
# ESSAY EXPORT
# ============================================

def _add_essay_docx(doc: Document, title: str, sections: Dict[str, List[str]], order: List[str]) -> None:
    """Add essay content to DOCX"""
    if title and title != "Improved Document":
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(12)
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(12)
        title_run.bold = False

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        for item in items:
            paragraph = doc.add_paragraph(item)
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 2


def _add_essay_pdf(story: list, title: str, sections: Dict[str, List[str]], order: List[str], pdf_profile: Dict) -> None:
    """Add essay content to PDF"""
    title_style = ParagraphStyle(
        "essay_title",
        fontName="Times-Roman",
        fontSize=pdf_profile.get("title_size", 14),
        alignment=1,
        textColor="#000000",
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "essay_body",
        fontName="Times-Roman",
        fontSize=pdf_profile.get("body_size", 12),
        leading=pdf_profile.get("body_leading", 24),
        firstLineIndent=0.5 * inch,
        spaceAfter=0,
        textColor="#000000",
    )

    if title and title != "Improved Document":
        story.append(Paragraph(_escape_html(title), title_style))

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        for item in items:
            story.append(Paragraph(_escape_html(item), body_style))


# ============================================
# PROPOSAL EXPORT
# ============================================

def _looks_like_subheading(item: str) -> bool:
    text = (item or "").strip()
    if not text:
        return False
    if len(text) > 80 or text.endswith((".", "!", "?")):
        return False
    words = text.split()
    if len(words) > 8:
        return False
    title_case_words = [
        word for word in words
        if regex.fullmatch(r"[A-Z][A-Za-z'`/&-]*", word.strip(",:;()"))
    ]
    uppercase_chars = sum(1 for char in text if char.isupper())
    alpha_chars = sum(1 for char in text if char.isalpha())
    return (
        bool(regex.match(r"^\d+(?:\.\d+)*\b", text))
        or (alpha_chars and uppercase_chars / alpha_chars >= 0.7)
        or (len(title_case_words) / max(1, len(words)) >= 0.75)
    )


def _is_numbered_subsection_line(item: str) -> bool:
    text = (item or "").strip()
    return bool(regex.match(r"^\d+\.\d+\b", text))

def _add_proposal_docx(doc: Document, title: str, contact: str, sections: Dict[str, List[str]], order: List[str]) -> None:
    """Add proposal content to DOCX"""
    accent = RGBColor(29, 68, 126)
    cover_page_kept = _add_cover_page_docx(doc, title, sections.get("Document Details", []), accent)

    if not cover_page_kept:
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(6)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = accent

        if contact:
            subtitle_paragraph = doc.add_paragraph()
            subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_paragraph.paragraph_format.space_after = Pt(10)
            subtitle_run = subtitle_paragraph.add_run(contact)
            subtitle_run.font.size = Pt(10)
            subtitle_run.font.color.rgb = RGBColor(84, 97, 120)

        divider = doc.add_paragraph()
        divider.paragraph_format.space_after = Pt(12)
        divider_run = divider.add_run(" ")
        _shade_paragraph(divider, fill="DCE8F8")
        divider_run.font.size = Pt(1)
        doc.add_page_break()

    first_written = False
    visible_sections = [section_name for section_name in order if section_name != "Document Details"]
    for index, section_name in enumerate(visible_sections, start=1):
        items = sections.get(section_name, [])
        if not items:
            continue

        if first_written:
            doc.add_page_break()

        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(8)
        heading_run = heading.add_run(f"{index}. {section_name}")
        heading_run.bold = True
        heading_run.font.size = Pt(15)
        heading_run.font.color.rgb = RGBColor(32, 74, 135)

        bullet_friendly = section_name in {"Objectives", "Deliverables", "Qualifications or Team"}
        subsection_count = 0
        for item in items:
            if bullet_friendly:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(5)
                paragraph.paragraph_format.left_indent = Inches(0.25)
                bullet_run = paragraph.add_run("\u2022 ")
                bullet_run.font.color.rgb = RGBColor(32, 74, 135)
                text_run = paragraph.add_run(item)
                text_run.font.size = Pt(10.8)
            elif _is_numbered_subsection_line(item):
                if subsection_count > 0:
                    _add_horizontal_rule_docx(doc, color="D7DFED")
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.25
                subsection_count += 1
            elif _looks_like_subheading(item):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                run = paragraph.add_run(item)
                run.bold = True
                run.font.size = Pt(11.8)
                run.font.color.rgb = RGBColor(55, 74, 107)
            else:
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.first_line_indent = Inches(0.3)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.25

        first_written = True


def _add_proposal_pdf(story: list, title: str, contact: str, sections: Dict[str, List[str]], order: List[str], pdf_profile: Dict) -> None:
    """Add proposal content to PDF"""
    title_style = ParagraphStyle(
        "proposal_title",
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("title_size", 18),
        alignment=1,
        textColor="#1D447E",
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "proposal_subtitle",
        fontName="Helvetica",
        fontSize=10,
        alignment=1,
        textColor="#556178",
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "proposal_heading",
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("heading_size", 12.5),
        textColor="#1D447E",
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "proposal_body",
        fontName="Helvetica",
        fontSize=pdf_profile.get("body_size", 10.8),
        leading=pdf_profile.get("body_leading", 15),
        spaceAfter=4,
    )

    story.append(Paragraph(_escape_html(title), title_style))
    if contact:
        story.append(Paragraph(_escape_html(contact), subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color="#C7D6EA", spaceAfter=8))

    for index, section_name in enumerate(order, start=1):
        items = sections.get(section_name, [])
        if not items:
            continue

        story.append(Paragraph(_escape_html(f"{index}. {section_name}"), heading_style))
        bullet_friendly = section_name in {"Objectives", "Deliverables", "Qualifications or Team"}

        for item in items:
            if bullet_friendly:
                story.append(Paragraph(_escape_html("\u2022 " + item), body_style))
            else:
                story.append(Paragraph(_escape_html(item), body_style))


# ============================================
# REPORT EXPORT
# ============================================

def _add_report_docx(doc: Document, title: str, contact: str, sections: Dict[str, List[str]], order: List[str]) -> None:
    """Add report content to DOCX"""
    accent = RGBColor(24, 42, 77)
    cover_page_kept = _add_cover_page_docx(doc, title, sections.get("Document Details", []), accent)

    if not cover_page_kept:
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(8)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = accent

        if contact:
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_paragraph.paragraph_format.space_after = Pt(10)
            meta_run = meta_paragraph.add_run(contact)
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(86, 98, 120)

        divider = doc.add_paragraph()
        divider.paragraph_format.space_after = Pt(14)
        divider_run = divider.add_run(" ")
        _shade_paragraph(divider, fill="DCE7F6")
        divider_run.font.size = Pt(1)
        doc.add_page_break()

    first_written = False
    visible_sections = [section_name for section_name in order if section_name != "Document Details"]
    for index, section_name in enumerate(visible_sections, start=1):
        items = sections.get(section_name, [])
        if not items:
            continue

        if first_written:
            doc.add_page_break()

        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(8)
        heading_run = heading.add_run(f"{index}.0 {section_name}")
        heading_run.bold = True
        heading_run.font.size = Pt(15)
        heading_run.font.color.rgb = RGBColor(28, 67, 124)

        if section_name == "Executive Summary":
            for item in items:
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.first_line_indent = Inches(0.35)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.3
                _shade_paragraph(paragraph, fill="F1F4F8")
            continue

        bullet_friendly = section_name in {"Recommendations", "Findings"}
        subsection_count = 0
        for item in items:
            if bullet_friendly:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(5)
                paragraph.paragraph_format.left_indent = Inches(0.25)
                bullet_run = paragraph.add_run("\u2022 ")
                bullet_run.font.color.rgb = RGBColor(28, 67, 124)
                text_run = paragraph.add_run(item)
                text_run.font.size = Pt(10.8)
            elif _is_numbered_subsection_line(item):
                if subsection_count > 0:
                    _add_horizontal_rule_docx(doc, color="D7DFED")
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.3
                subsection_count += 1
            elif _looks_like_subheading(item):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                run = paragraph.add_run(item)
                run.bold = True
                run.font.size = Pt(11.8)
                run.font.color.rgb = RGBColor(46, 69, 104)
            else:
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.first_line_indent = Inches(0.35)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.3
        first_written = True


def _add_report_pdf(story: list, title: str, contact: str, sections: Dict[str, List[str]], order: List[str], pdf_profile: Dict) -> None:
    """Add report content to PDF"""
    title_style = ParagraphStyle(
        "report_title",
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("title_size", 19.5),
        alignment=1,
        textColor="#1A2D52",
        spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        "report_meta",
        fontName="Helvetica",
        fontSize=10,
        alignment=1,
        textColor="#556178",
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "report_heading",
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("heading_size", 12),
        textColor="#1D447E",
        spaceBefore=9,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "report_body",
        fontName="Helvetica",
        fontSize=pdf_profile.get("body_size", 10.5),
        leading=pdf_profile.get("body_leading", 15),
        spaceAfter=4,
    )
    summary_style = ParagraphStyle(
        "report_summary",
        parent=body_style,
        backColor="#F1F4F8",
        borderPadding=6,
    )

    story.append(Paragraph(_escape_html(title), title_style))
    if contact:
        story.append(Paragraph(_escape_html(contact), meta_style))
    story.append(HRFlowable(width="100%", thickness=1, color="#C7D6EA", spaceAfter=8))

    for index, section_name in enumerate(order, start=1):
        items = sections.get(section_name, [])
        if not items:
            continue

        story.append(Paragraph(_escape_html(f"{index}.0 {section_name}"), heading_style))
        bullet_friendly = section_name in {"Recommendations", "Findings"}

        for item in items:
            if section_name == "Executive Summary":
                story.append(Paragraph(_escape_html(item), summary_style))
            elif bullet_friendly:
                story.append(Paragraph(_escape_html("• " + item), body_style))
            else:
                story.append(Paragraph(_escape_html(item), body_style))


# ============================================
# THESIS EXPORT
# ============================================

def _add_thesis_docx(doc: Document, title: str, contact: str, sections: Dict[str, List[str]], order: List[str]) -> None:
    """Add thesis content to DOCX"""
    cover_page_kept = _add_cover_page_docx(doc, title, sections.get("Document Details", []), RGBColor(0, 0, 0))

    if not cover_page_kept:
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(18)
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0, 0, 0)

        if contact:
            meta_paragraph = doc.add_paragraph()
            meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_paragraph.paragraph_format.space_after = Pt(14)
            meta_run = meta_paragraph.add_run(contact)
            meta_run.font.size = Pt(11)
        doc.add_page_break()

    visible_sections = [section_name for section_name in order if section_name != "Document Details"]
    for index, section_name in enumerate(visible_sections, start=1):
        items = sections.get(section_name, [])
        if not items:
            continue

        if index > 1:
            doc.add_page_break()

        chapter = doc.add_paragraph()
        chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter.paragraph_format.space_before = Pt(90)
        chapter.paragraph_format.space_after = Pt(8)
        chapter_run = chapter.add_run(f"CHAPTER {index}")
        chapter_run.bold = True
        chapter_run.font.size = Pt(14)

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_after = Pt(16)
        heading_run = heading.add_run(section_name.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(14)

        for item in items:
            paragraph = doc.add_paragraph(item)
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.5


def _add_thesis_pdf(story: list, title: str, contact: str, sections: Dict[str, List[str]], order: List[str], pdf_profile: Dict) -> None:
    """Add thesis content to PDF"""
    title_style = ParagraphStyle(
        "thesis_title",
        fontName="Times-Bold",
        fontSize=pdf_profile.get("title_size", 20),
        alignment=1,
        textColor="#000000",
        spaceAfter=16,
    )
    meta_style = ParagraphStyle(
        "thesis_meta",
        fontName="Times-Roman",
        fontSize=11,
        alignment=1,
        textColor="#000000",
        spaceAfter=14,
    )
    chapter_style = ParagraphStyle(
        "thesis_chapter",
        fontName="Times-Bold",
        fontSize=14,
        alignment=1,
        textColor="#000000",
        spaceBefore=70,
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "thesis_heading",
        fontName="Times-Bold",
        fontSize=pdf_profile.get("heading_size", 13),
        alignment=1,
        textColor="#000000",
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "thesis_body",
        fontName="Times-Roman",
        fontSize=pdf_profile.get("body_size", 12),
        leading=pdf_profile.get("body_leading", 18),
        firstLineIndent=0.5 * inch,
        spaceAfter=6,
        textColor="#000000",
    )

    story.append(Paragraph(_escape_html(title), title_style))
    if contact:
        story.append(Paragraph(_escape_html(contact), meta_style))

    for index, section_name in enumerate(order, start=1):
        items = sections.get(section_name, [])
        if not items:
            continue

        if index > 1:
            story.append(PageBreak())

        story.append(Paragraph(_escape_html(f"CHAPTER {index}"), chapter_style))
        story.append(Paragraph(_escape_html(section_name.upper()), heading_style))
        for item in items:
            story.append(Paragraph(_escape_html(item), body_style))


# ============================================
# RESUME/CV EXPORT
# ============================================

def _add_resume_docx(doc: Document, title: str, contact: str, sections: Dict[str, List[str]], order: List[str]) -> None:
    """Add resume content to DOCX"""
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = Pt(5)
    name_run = name_paragraph.add_run(title)
    name_run.bold = True
    name_run.font.size = Pt(17)
    name_run.font.color.rgb = RGBColor(24, 41, 77)

    if contact:
        contact_lines = _split_resume_contact_lines(contact)
        for index, line in enumerate(contact_lines):
            contact_paragraph = doc.add_paragraph()
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_paragraph.paragraph_format.space_after = Pt(2 if index < len(contact_lines) - 1 else 8)
            contact_run = contact_paragraph.add_run(line)
            contact_run.font.size = Pt(9.6)
            contact_run.font.color.rgb = RGBColor(70, 78, 96)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(8)
    divider_run = divider.add_run(" ")
    _shade_paragraph(divider, fill="D8E4F6")
    divider_run.font.size = Pt(1)

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        heading_run = heading.add_run(section_name.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(12.5)
        heading_run.font.color.rgb = RGBColor(26, 61, 111)
        _shade_paragraph(heading, fill="EEF3FB")

        if section_name == "Professional Summary":
            for item in items:
                paragraph = doc.add_paragraph(item)
                paragraph.paragraph_format.space_after = Pt(4)
            continue

        for item in items:
            line_type = _resume_line_type(section_name, item)
            if section_name == "Education" and line_type == "header" and not _education_item_is_institution(item):
                line_type = "detail"
            if section_name == "Experience" and _experience_item_should_be_header(item):
                line_type = "header"
            _add_resume_entry_line_docx(doc, item, line_type)


def _add_resume_pdf(story: list, title: str, contact: str, sections: Dict[str, List[str]], order: List[str], pdf_profile: Dict) -> None:
    """Add resume content to PDF"""
    title_style = ParagraphStyle(
        "resume_title",
        fontName="Helvetica-Bold",
        fontSize=min(pdf_profile.get("title_size", 26), 24),
        alignment=1,
        textColor="#1A2D52",
        leading=28,
        spaceAfter=6,
    )
    contact_style = ParagraphStyle(
        "resume_contact",
        fontName="Helvetica",
        fontSize=9.2,
        alignment=1,
        textColor="#4B5870",
        leading=11.5,
        spaceAfter=2,
    )
    heading_style = ParagraphStyle(
        "resume_heading",
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("heading_size", 12),
        textColor="#1E4A86",
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "resume_body",
        fontName="Helvetica",
        fontSize=pdf_profile.get("body_size", 10),
        leading=pdf_profile.get("body_leading", 14.5),
        spaceAfter=2.5,
    )

    story.append(Paragraph(_escape_html(title), title_style))
    if contact:
        for line in _split_resume_contact_lines(contact):
            story.append(Paragraph(_escape_html(line), contact_style))
        story.append(Spacer(1, 0.04 * inch))
    story.append(HRFlowable(width="100%", thickness=0.8, color="#C7D6EA", spaceAfter=8))

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        story.append(Paragraph(_escape_html(section_name.upper()), heading_style))
        if section_name == "Professional Summary":
            for item in items:
                story.append(Paragraph(_escape_html(item), body_style))
            continue

        for item in items:
            line_type = _resume_line_type(section_name, item)
            if section_name == "Education" and line_type == "header" and not _education_item_is_institution(item):
                line_type = "detail"
            if section_name == "Experience" and _experience_item_should_be_header(item):
                line_type = "header"
            _add_resume_entry_line_pdf(story, item, line_type, body_style)

        story.append(Spacer(1, 0.03 * inch))


# ============================================
# MAIN EXPORT FUNCTIONS
# ============================================

def export_docx(
    title: str,
    contact: str,
    sections: Dict[str, List[str]],
    order: List[str],
    output_path: Path,
    doc_type: str = "General Document",
) -> None:
    """
    Export document to DOCX format.
    
    Args:
        title: Document title
        contact: Contact information string
        sections: Dictionary of section names to content lists
        order: Ordered list of section names
        output_path: Path where to save the DOCX file
        doc_type: Type of document (Resume, Cover Letter, Thesis, etc.)
    """
    profile = get_format_profile(doc_type)
    docx_profile = profile.get("docx", {})
    heading_size = docx_profile.get("heading_size", 13)
    heading_space_before = docx_profile.get("heading_space_before", 10)
    heading_space_after = docx_profile.get("heading_space_after", 4)
    body_space_after = docx_profile.get("body_space_after", 3)
    page_break_before = set(profile.get("page_break_before", []))

    doc = Document()
    _setup_docx(doc, profile)

    # Handle specific document types
    if doc_type == "Cover Letter":
        _add_cover_letter_docx(doc, title, contact, sections, order, body_space_after)
        doc.save(output_path)
        return
    
    if doc_type == "Essay":
        _add_essay_docx(doc, title, sections, order)
        doc.save(output_path)
        return
    
    if doc_type == "Proposal":
        _add_proposal_docx(doc, title, contact, sections, order)
        doc.save(output_path)
        return
    
    if doc_type == "Thesis":
        _add_thesis_docx(doc, title, contact, sections, order)
        doc.save(output_path)
        return
    
    if doc_type == "Report":
        _add_report_docx(doc, title, contact, sections, order)
        doc.save(output_path)
        return
    
    if doc_type in {"Resume", "CV"}:
        _add_resume_docx(doc, title, contact, sections, order)
        doc.save(output_path)
        return

    # General document fallback
    _add_general_header_docx(doc, title, contact)

    first_section_written = False

    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        if _needs_page_break(section_name, page_break_before, first_section_written):
            doc.add_page_break()

        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(heading_space_before)
        heading.paragraph_format.space_after = Pt(heading_space_after)
        _shade_paragraph(heading, fill="EAF1FF")
        heading_run = heading.add_run(section_name.upper())
        heading_run.bold = True
        heading_run.font.size = Pt(heading_size)
        heading_run.font.color.rgb = RGBColor(18, 57, 102)

        if section_name in PARAGRAPH_SECTIONS:
            for item in items:
                p = doc.add_paragraph(item)
                p.paragraph_format.space_after = Pt(body_space_after)
        else:
            for item in items:
                p = doc.add_paragraph(f"- {item}")
                p.paragraph_format.space_after = Pt(max(1.5, body_space_after - 1))

        first_section_written = True

    doc.save(output_path)


def export_pdf(
    title: str,
    contact: str,
    sections: Dict[str, List[str]],
    order: List[str],
    output_path: Path,
    doc_type: str = "General Document",
) -> None:
    """
    Export document to PDF format.
    
    Args:
        title: Document title
        contact: Contact information string
        sections: Dictionary of section names to content lists
        order: Ordered list of section names
        output_path: Path where to save the PDF file
        doc_type: Type of document (Resume, Cover Letter, Thesis, etc.)
    """
    profile = get_format_profile(doc_type)
    pdf_profile = profile.get("pdf", {})
    page_break_before = set(profile.get("page_break_before", []))

    styles = getSampleStyleSheet()
    resume_mode = doc_type in {"Resume", "CV"}

    title_style = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("title_size", 20),
        alignment=0 if resume_mode else 1,
        textColor="#1A2D52",
        spaceAfter=5,
    )
    contact_style = ParagraphStyle(
        "contact",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.4,
        alignment=0,
        textColor="#42526E",
        spaceAfter=7,
    )
    heading_style = ParagraphStyle(
        "heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=pdf_profile.get("heading_size", 12),
        textColor="#1D447E",
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=pdf_profile.get("body_size", 10),
        leading=pdf_profile.get("body_leading", 14.5),
        spaceAfter=3,
    )

    story = []

    # Handle specific document types
    if doc_type == "Cover Letter":
        _add_cover_letter_pdf(story, title, contact, sections, order, body_style)
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=1.0 * inch,
            rightMargin=1.0 * inch,
            topMargin=1.0 * inch,
            bottomMargin=1.0 * inch,
        )
        pdf.build(story)
        return
    
    if doc_type == "Essay":
        _add_essay_pdf(story, title, sections, order, pdf_profile)
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=1.0 * inch,
            rightMargin=1.0 * inch,
            topMargin=1.0 * inch,
            bottomMargin=1.0 * inch,
        )
        pdf.build(story)
        return
    
    if doc_type == "Proposal":
        _add_proposal_pdf(story, title, contact, sections, order, pdf_profile)
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=0.9 * inch,
            rightMargin=0.9 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch,
        )
        pdf.build(story)
        return
    
    if doc_type == "Thesis":
        _add_thesis_pdf(story, title, contact, sections, order, pdf_profile)
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=1.0 * inch,
            rightMargin=1.0 * inch,
            topMargin=1.0 * inch,
            bottomMargin=1.0 * inch,
        )
        pdf.build(story)
        return
    
    if doc_type == "Report":
        _add_report_pdf(story, title, contact, sections, order, pdf_profile)
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=0.9 * inch,
            rightMargin=0.9 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch,
        )
        pdf.build(story)
        return
    
    if doc_type in {"Resume", "CV"}:
        _add_resume_pdf(story, title, contact, sections, order, pdf_profile)
        pdf = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.6 * inch,
            bottomMargin=0.6 * inch,
        )
        pdf.build(story)
        return

    # General document fallback
    story.append(Paragraph(_escape_html(title), title_style))
    story.append(HRFlowable(width="100%", thickness=1.2, color="#7BA3E6", spaceAfter=8))

    if contact:
        story.append(Paragraph(_escape_html(contact), contact_style))

    story.append(Spacer(1, 0.04 * inch))

    first_section_written = False
    for section_name in order:
        items = sections.get(section_name, [])
        if not items:
            continue

        if _needs_page_break(section_name, page_break_before, first_section_written):
            story.append(PageBreak())

        if resume_mode and section_name == "Professional Summary":
            continue

        story.append(Paragraph(_escape_html(section_name.upper()), heading_style))

        if section_name in PARAGRAPH_SECTIONS:
            for item in items:
                story.append(Paragraph(_escape_html(item), body_style))
        else:
            for item in items:
                story.append(Paragraph(_escape_html("- " + item), body_style))

        story.append(Spacer(1, 0.05 * inch))
        first_section_written
